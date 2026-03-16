# -*- coding: utf-8 -*-
"""
Xuat Word tu file Excel da tach theo tung nguoi.

Muc tieu:
- Tao title "CHI TIET CHAM CONG"
- Hien thi bang cham cong cua tung nguoi de in
- Ho tro ca .xls va .xlsx
"""

import os
import re
import unicodedata
from datetime import date, datetime, time, timedelta
from typing import List, Optional

import xlrd
from openpyxl import load_workbook
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def _normalize(text: str) -> str:
    text = unicodedata.normalize('NFD', str(text).strip().lower())
    text = ''.join(c for c in text if unicodedata.category(c) != 'Mn')
    return re.sub(r'\s+', ' ', text)


def _format_number(value) -> str:
    try:
        number = float(value)
    except Exception:
        return str(value).strip()

    if abs(number - round(number)) < 1e-9:
        return str(int(round(number)))
    return f"{number:.2f}".rstrip('0').rstrip('.')


def _format_date_value(dt: date) -> str:
    return f"{dt.month}/{dt.day}/{dt.year}"


def _format_time_value(value: time) -> str:
    return f"{value.hour:02d}:{value.minute:02d}"


def _format_timedelta_value(delta: timedelta) -> str:
    total_seconds = int(delta.total_seconds())
    if total_seconds < 0:
        total_seconds = 0
    hours = total_seconds // 3600
    minutes = (total_seconds % 3600) // 60
    return f"{hours:02d}:{minutes:02d}"


def _format_display_value(value) -> str:
    if value is None:
        return ''
    if isinstance(value, datetime):
        return _format_date_value(value.date())
    if isinstance(value, date):
        return _format_date_value(value)
    if isinstance(value, time):
        return _format_time_value(value)
    if isinstance(value, timedelta):
        return _format_timedelta_value(value)
    if isinstance(value, bool):
        return 'TRUE' if value else 'FALSE'
    if isinstance(value, (int, float)):
        return _format_number(value)
    return str(value).strip()


def _format_xlrd_cell(cell, datemode: int) -> str:
    if cell.ctype in (xlrd.XL_CELL_EMPTY, xlrd.XL_CELL_BLANK):
        return ''

    if cell.ctype == xlrd.XL_CELL_DATE:
        year, month, day, hour, minute, second = xlrd.xldate_as_tuple(
            cell.value,
            datemode
        )
        if year == 0 and month == 0 and day == 0:
            return f"{hour:02d}:{minute:02d}"
        if hour or minute or second:
            return f"{month}/{day}/{year} {hour:02d}:{minute:02d}"
        return f"{month}/{day}/{year}"

    if cell.ctype == xlrd.XL_CELL_NUMBER:
        return _format_number(cell.value)

    if cell.ctype == xlrd.XL_CELL_BOOLEAN:
        return 'TRUE' if cell.value else 'FALSE'

    return str(cell.value).strip()


def _read_rows(path: str) -> List[List[str]]:
    ext = os.path.splitext(path)[1].lower()
    rows: List[List[str]] = []

    if ext == '.xlsx':
        workbook = load_workbook(path, data_only=True)
        worksheet = workbook.active
        for row in worksheet.iter_rows():
            rows.append([_format_display_value(cell.value) for cell in row])
    else:
        workbook = xlrd.open_workbook(path)
        worksheet = workbook.sheet_by_index(0)
        for row_index in range(worksheet.nrows):
            rows.append([
                _format_xlrd_cell(worksheet.cell(row_index, col_index), workbook.datemode)
                for col_index in range(worksheet.ncols)
            ])

    return rows


class ExcelListWordExporter:
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def _find_header_row(self, rows: List[List[str]]) -> Optional[int]:
        for row_index, row in enumerate(rows[:50]):
            line = " ".join(_normalize(value) for value in row if value != '')
            if "ma nhan vien" in line and "ten nhan vien" in line and "ngay" in line:
                return row_index
        return None

    def _find_title(self, rows: List[List[str]]) -> str:
        for row in rows[:10]:
            line = " ".join(str(value).strip() for value in row if value)
            if "chi tiet cham cong" in _normalize(line):
                return "CHI TIẾT CHẤM CÔNG"
        return "CHI TIẾT CHẤM CÔNG"

    def _find_last_used_col(self, rows: List[List[str]], header_idx: int) -> int:
        last_col = 0
        for row in rows[header_idx:]:
            for index, value in enumerate(row):
                if str(value).strip() != '':
                    last_col = max(last_col, index)
        return last_col

    def _column_width(self, header_text: str) -> Cm:
        norm = _normalize(header_text)
        fixed_widths = {
            'stt': 1.0,
            'ma nhan vien': 2.3,
            'ten nhan vien': 4.8,
            'phong ban': 2.4,
            'ngay': 2.0,
            'thu': 1.3,
            'gio vao': 1.6,
            'gio ra': 1.6,
            'tre': 1.0,
            'som': 1.0,
            'cong': 1.0,
            'tong gio': 1.8,
            'tang ca': 1.7,
            'tong toan bo': 2.2,
            'ca': 1.2,
        }
        if norm in fixed_widths:
            return Cm(fixed_widths[norm])
        if 'ten' in norm:
            return Cm(4.2)
        if 'gio' in norm:
            return Cm(1.8)
        if 'tong' in norm:
            return Cm(1.9)
        return Cm(1.6)

    def _cell_alignment(self, header_text: str):
        norm = _normalize(header_text)
        if norm in ('ten nhan vien', 'phong ban'):
            return WD_ALIGN_PARAGRAPH.LEFT
        return WD_ALIGN_PARAGRAPH.CENTER

    def _set_run_font(self, run, size: float, bold: bool = False):
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'
        run_properties = run._element.get_or_add_rPr()
        run_fonts = run_properties.rFonts
        if run_fonts is None:
            run_fonts = OxmlElement('w:rFonts')
            run_properties.append(run_fonts)
        run_fonts.set(qn('w:ascii'), 'Times New Roman')
        run_fonts.set(qn('w:hAnsi'), 'Times New Roman')
        run_fonts.set(qn('w:eastAsia'), 'Times New Roman')

    def _shade_cell(self, cell, fill: str):
        cell_properties = cell._tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), fill)
        cell_properties.append(shading)

    def _set_fixed_layout(self, table):
        table_properties = table._tbl.tblPr
        layout = OxmlElement('w:tblLayout')
        layout.set(qn('w:type'), 'fixed')
        table_properties.append(layout)

    def _repeat_header(self, row):
        row_properties = row._tr.get_or_add_trPr()
        table_header = OxmlElement('w:tblHeader')
        table_header.set(qn('w:val'), 'true')
        row_properties.append(table_header)

    def export_from_excel(self, excel_path: str) -> Optional[str]:
        rows = _read_rows(excel_path)
        if not rows:
            return None

        header_idx = self._find_header_row(rows)
        if header_idx is None:
            return None

        title = self._find_title(rows)
        last_col = self._find_last_used_col(rows, header_idx)
        header = [str(value).strip() for value in rows[header_idx][:last_col + 1]]

        data_rows = []
        for raw_row in rows[header_idx + 1:]:
            row = (raw_row + [''] * (last_col + 1))[:last_col + 1]
            if all(str(value).strip() == '' for value in row):
                continue
            data_rows.append([str(value).strip() for value in row])

        document = Document()
        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)

        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        self._set_run_font(title_run, 18, bold=True)

        document.add_paragraph()

        table = document.add_table(rows=1, cols=len(header))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_fixed_layout(table)

        header_row = table.rows[0]
        self._repeat_header(header_row)

        for index, text in enumerate(header):
            cell = header_row.cells[index]
            cell.width = self._column_width(text)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self._shade_cell(cell, 'D9D9D9')
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(text)
            self._set_run_font(run, 10, bold=True)

        for row_data in data_rows:
            row = table.add_row()
            for index, value in enumerate(row_data):
                cell = row.cells[index]
                cell.width = self._column_width(header[index])
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph = cell.paragraphs[0]
                paragraph.alignment = self._cell_alignment(header[index])
                run = paragraph.add_run(value)
                self._set_run_font(run, 9.5)

        base_name = os.path.splitext(os.path.basename(excel_path))[0]
        output_path = os.path.join(self.output_dir, f"{base_name}.docx")
        document.save(output_path)
        return output_path
