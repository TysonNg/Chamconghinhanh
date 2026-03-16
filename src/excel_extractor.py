# -*- coding: utf-8 -*-
"""
Module xá»­ lÃ½ file Excel cháº¥m cÃ´ng (Ä‘á»‹nh dáº¡ng Ä‘áº·c biá»‡t vá»›i cá»™t lÃ  ngÃ y)
TrÃ­ch xuáº¥t thÃ´ng tin tá»«ng ngÆ°á»i vÃ  xuáº¥t ra file Word riÃªng
"""

import os
import re
import unicodedata
import xlrd
from datetime import datetime, date
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Dict, Optional, Tuple


def set_cell_border(cell, border_color="000000"):
    """Set cell border for Word table cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


class ExcelChamCongExtractor:
    """
    Äá»c file Excel cháº¥m cÃ´ng dáº¡ng pivot (cá»™t = ngÃ y, má»—i ngÆ°á»i 6 dÃ²ng)
    vÃ  tÃ¡ch thÃ nh dá»¯ liá»‡u tá»«ng ngÆ°á»i.
    
    Cáº¥u trÃºc má»—i khá»‘i ngÆ°á»i:
      Row N+0: ID:xxx TÃªn:yyy PhÃ²ng ban:zzz Ca:www
      Row N+1: 1  2  3  4 ... 16   (sá»‘ ngÃ y ná»­a Ä‘áº§u thÃ¡ng)
      Row N+2: giá» vÃ o\ngiá» ra ... (dá»¯ liá»‡u ná»­a Ä‘áº§u thÃ¡ng)
      Row N+3: 17 18 19 ... cuá»‘i  (sá»‘ ngÃ y ná»­a sau thÃ¡ng)
      Row N+4: giá» vÃ o\ngiá» ra ... (dá»¯ liá»‡u ná»­a sau thÃ¡ng)
      Row N+5: (trá»‘ng - phÃ¢n cÃ¡ch)
    """

    def __init__(self, excel_path: str):
        self.excel_path = excel_path
        self.wb = xlrd.open_workbook(excel_path)
        self.sheet = self.wb.sheet_by_index(0)
        self._persons_data = []  # List of dicts per person
        self._date_range = None
        self._header_map = None
        self._parse_all()

    def _normalize_text(self, text: str) -> str:
        """Normalize text for robust header matching (lowercase, remove accents, trim spaces)."""
        if not text:
            return ""
        text = str(text)
        text = unicodedata.normalize('NFD', text)
        text = ''.join(c for c in text if unicodedata.category(c) != 'Mn')
        text = re.sub(r'\s+', ' ', text.lower().strip())
        return text

    def _detect_list_header(self) -> Optional[Dict[str, int]]:
        """Detect list-format header row and return column map if found."""
        if self.sheet.nrows == 0:
            return None

        keywords = {
            'id': ['ma nhan vien', 'ma the', 'ma nv', 'id', 'manv', 'mathe'],
            'name': ['ho ten', 'ten nhan vien', 'ho va ten', 'ten'],
            'date': ['ngay', 'date', 'ngay cong', 'ngay lam'],
            'gio_vao': ['gio vao', 'check in', 'time in', 'vao', 'giovao'],
            'gio_ra': ['gio ra', 'check out', 'time out', 'ra', 'giora'],
        }

        best = None
        best_row = None

        max_rows = min(20, self.sheet.nrows)
        for r in range(max_rows):
            col_map = {}
            strength = 0
            for c in range(self.sheet.ncols):
                val = self.sheet.cell_value(r, c)
                norm = self._normalize_text(val)
                if not norm:
                    continue
                for key, keys in keywords.items():
                    for k in keys:
                        if k in norm:
                            if key not in col_map or len(k) > len(col_map[key][1]):
                                col_map[key] = (c, k)
            if col_map:
                # Count how many essential columns detected
                for k in ('id', 'name', 'date'):
                    if k in col_map:
                        strength += 1
                if strength >= 2:
                    if best is None or strength > best:
                        best = strength
                        best_row = col_map

        if not best_row:
            return None

        # Convert to simple map {field: col_idx}
        return {k: v[0] for k, v in best_row.items()}

    # ------------------------------------------------------------------
    # Internal parsers
    # ------------------------------------------------------------------

    def _parse_info_row(self, row_idx: int) -> Optional[Dict]:
        """Parse dÃ²ng thÃ´ng tin: 'ID:xxx TÃªn:yyy PhÃ²ng ban:zzz Ca:www'"""
        o_variants = "oòóỏõọôồốổỗộơờớởỡợ"
        e_variants = "eèéẻẽẹêềếểễệ"
        ten_pat = f"T[{e_variants}]n"
        phong_ban_pat = f"Ph[{o_variants}]ng\\s*ban"
        info_pat = re.compile(
            rf"ID[:\s]*([^\s]+).*?{ten_pat}[:\s]*(.+?)(?:{phong_ban_pat}|$)",
            re.IGNORECASE
        )

        for c in range(self.sheet.ncols):
            cell_val = str(self.sheet.cell_value(row_idx, c)).strip()
            if not cell_val:
                continue
            # Quick filter to avoid unnecessary regex
            norm = self._normalize_text(cell_val)
            if "id" not in norm or "ten" not in norm:
                continue
            m = info_pat.search(cell_val)
            if not m:
                continue
            emp_id = m.group(1).strip()
            name = m.group(2).strip()
            return {'id': emp_id, 'name': name, 'raw_info': cell_val}
        return None

    def _parse_days_and_times(self, day_row_idx: int, time_row_idx: int) -> Dict[int, Dict]:
        """
        Parse má»™t cáº·p (dÃ²ng ngÃ y, dÃ²ng giá») â†’ dict {ngÃ y: {gio_vao, gio_ra, vang}}
        """
        results = {}
        ncols = self.sheet.ncols
        for c in range(ncols):
            day_val = self.sheet.cell_value(day_row_idx, c)
            time_val = self.sheet.cell_value(time_row_idx, c)

            if day_val == '' or day_val is None:
                continue
            try:
                day_num = int(float(str(day_val).strip()))
            except ValueError:
                continue

            gio_vao = ''
            gio_ra = ''
            is_absent = True

            if time_val and str(time_val).strip():
                lines = [l.strip() for l in str(time_val).split('\n') if l.strip()]
                # First valid time = giá» vÃ o, last valid time = giá» ra
                time_pattern = re.compile(r'^\d{1,2}:\d{2}$')
                valid_times = [l for l in lines if time_pattern.match(l)]
                if valid_times:
                    gio_vao = valid_times[0]
                    gio_ra = valid_times[-1] if len(valid_times) > 1 else ''
                    is_absent = False  # cÃ³ Ã­t nháº¥t 1 giá»
                    # Náº¿u chá»‰ cÃ³ 1 giá» = thiáº¿u giá» ra
                    if len(valid_times) == 1:
                        is_absent = False  # cÃ³ máº·t nhÆ°ng thiáº¿u 1 cá»™t

            results[day_num] = {
                'gio_vao': gio_vao,
                'gio_ra': gio_ra,
                'is_absent': is_absent,
                'missing_checkout': bool(gio_vao and not gio_ra),
                'missing_checkin': bool(not gio_vao and gio_ra),
            }
        return results

    def _parse_all(self):
        """QuÃ©t toÃ n bá»™ sheet vÃ  á»§y quyá»n parse cho Ä‘á»‹nh dáº¡ng phÃ¹ há»£p"""
        # Detect list format by scanning header row
        header_map = self._detect_list_header()
        if header_map:
            self._header_map = header_map
            self._parse_list_format()
        else:
            self._parse_pivot_format()
        self._dedupe_persons_by_name()

    def _parse_pivot_format(self):
        """Parse Ä‘á»‹nh dáº¡ng Pivot (ngang): má»—i nhÃ¢n sá»± gá»“m khá»‘i 6 dÃ²ng"""
        # Get date range from row 1
        date_info = str(self.sheet.cell_value(1, 0))
        m = re.search(r'(\d{4}-\d{1,2}-\d{1,2})~(\d{4}-\d{1,2}-\d{1,2})', date_info)
        if m:
            self._date_range = (m.group(1), m.group(2))
            # Extract month/year from start date
            start_dt = datetime.strptime(m.group(1), '%Y-%m-%d')
            self._month = start_dt.month
            self._year = start_dt.year
        else:
            now = datetime.now()
            self._month = now.month
            self._year = now.year

        nrows = self.sheet.nrows
        r = 0
        while r < nrows:
            info = self._parse_info_row(r)
            if info and r + 4 < nrows:
                # Parse first half-month (rows r+1=days, r+2=times)
                days_1 = self._parse_days_and_times(r + 1, r + 2)
                # Parse second half-month (rows r+3=days, r+4=times)
                days_2 = self._parse_days_and_times(r + 3, r + 4)
                # Merge
                all_days = {**days_1, **days_2}
                # Build daily records list
                records = []
                for day_num in sorted(all_days.keys()):
                    try:
                        rec_date = date(self._year, self._month, day_num)
                    except ValueError:
                        continue
                    day_data = all_days[day_num]
                    records.append({
                        'day': day_num,
                        'date': rec_date.strftime('%d/%m/%Y'),
                        'weekday': ['CN', 'Hai', 'Ba', 'TÆ°', 'NÄƒm', 'SÃ¡u', 'Báº£y'][rec_date.weekday() % 7
                                    if rec_date.weekday() != 6 else 0],
                        'gio_vao': day_data['gio_vao'],
                        'gio_ra': day_data['gio_ra'],
                        'is_absent': day_data['is_absent'],
                        'missing_checkout': day_data['missing_checkout'],
                        'missing_checkin': day_data['missing_checkin'],
                    })

                self._persons_data.append({
                    'id': info['id'],
                    'name': info['name'],
                    'month': self._month,
                    'year': self._year,
                    'records': records,
                })
                r += 5  # Skip person block
                continue
            r += 1

    def _parse_list_format(self):
        """Parse Ä‘á»‹nh dáº¡ng List dá»c: Má»—i dÃ²ng thá»ƒ hiá»‡n 1 ca lÃ m viá»‡c cá»§a nhÃ¢n sá»± trong ngÃ y nháº¥t Ä‘á»‹nh"""
        persons_dict = {}
        
        # Check title row for month/year or use fallback
        self._year, self._month = datetime.now().year, datetime.now().month
        header_map = self._header_map or {}
        id_col = header_map.get('id', 1)
        name_col = header_map.get('name', 2)
        date_col = header_map.get('date', 4)
        gio_vao_col = header_map.get('gio_vao', 6)
        gio_ra_col = header_map.get('gio_ra', 7)

        # Start after header row if detected
        start_row = 0
        if header_map:
            for r in range(min(20, self.sheet.nrows)):
                row_norm = [self._normalize_text(self.sheet.cell_value(r, c)) for c in range(self.sheet.ncols)]
                if any('ma nhan vien' in v or 'ma the' in v for v in row_norm):
                    start_row = r + 1
                    break
        if start_row == 0:
            start_row = 3

        for r in range(start_row, self.sheet.nrows):
            emp_id = str(self.sheet.cell_value(r, id_col)).strip()
            if not emp_id or self._normalize_text(emp_id) in ('ma nhan vien', 'ma the', 'ma nv', 'id'):
                continue

            name = str(self.sheet.cell_value(r, name_col)).strip()
            day_val = self.sheet.cell_value(r, date_col)
            if not day_val:
                continue

            # Convert Excel date
            try:
                if isinstance(day_val, (int, float)):
                    dt = xlrd.xldate_as_tuple(float(day_val), self.wb.datemode)
                    rec_date = date(dt[0], dt[1], dt[2])
                    self._year, self._month = dt[0], dt[1]
                else:
                    day_text = str(day_val).strip()
                    # Accept common date formats like 2/1/2026 or 02-01-2026
                    rec_date = None
                    for fmt in ("%d/%m/%Y", "%m/%d/%Y", "%d-%m-%Y", "%m-%d-%Y"):
                        try:
                            rec_date = datetime.strptime(day_text, fmt).date()
                            break
                        except Exception:
                            pass
                    if rec_date is None:
                        continue
                    self._year, self._month = rec_date.year, rec_date.month
            except Exception:
                continue

            # Get in/out values
            gio_vao = str(self.sheet.cell_value(r, gio_vao_col)).strip()
            gio_ra = str(self.sheet.cell_value(r, gio_ra_col)).strip()

            if self.sheet.cell_type(r, gio_vao_col) in (xlrd.XL_CELL_DATE, xlrd.XL_CELL_NUMBER) and gio_vao:
                try:
                    vt = xlrd.xldate_as_tuple(self.sheet.cell_value(r, gio_vao_col), self.wb.datemode)
                    gio_vao = f"{vt[3]:02d}:{vt[4]:02d}"
                except Exception:
                    pass

            if self.sheet.cell_type(r, gio_ra_col) in (xlrd.XL_CELL_DATE, xlrd.XL_CELL_NUMBER) and gio_ra:
                try:
                    rt = xlrd.xldate_as_tuple(self.sheet.cell_value(r, gio_ra_col), self.wb.datemode)
                    gio_ra = f"{rt[3]:02d}:{rt[4]:02d}"
                except Exception:
                    pass

            if emp_id not in persons_dict:
                persons_dict[emp_id] = {
                    'id': emp_id,
                    'name': name,
                    'month': self._month,
                    'year': self._year,
                    'records_dict': {}
                }

            day_num = rec_date.day
            day_key = rec_date.strftime('%Y%m%d')
            
            p_records = persons_dict[emp_id]['records_dict']
            if day_key not in p_records:
                p_records[day_key] = {
                    'day': day_num,
                    'date': rec_date.strftime('%d/%m/%Y'),
                    'weekday': ['CN', 'Hai', 'Ba', 'TÆ°', 'NÄƒm', 'SÃ¡u', 'Báº£y'][rec_date.weekday() % 7 if rec_date.weekday() != 6 else 0],
                    'vao_list': [],
                    'ra_list': []
                }

            # Ná»‘i cÃ¡c giá» ca láº¡i
            if gio_vao: p_records[day_key]['vao_list'].append(gio_vao)
            if gio_ra: p_records[day_key]['ra_list'].append(gio_ra)

        for emp_id, pdata in persons_dict.items():
            records = []
            for d_key in sorted(pdata['records_dict'].keys()):
                r = pdata['records_dict'][d_key]
                vaos = r['vao_list']
                ras = r['ra_list']
                
                is_absent = len(vaos) == 0 and len(ras) == 0
                missing_checkout = bool(len(vaos) > len(ras))
                missing_checkin = bool(len(ras) > len(vaos))
                
                records.append({
                    'day': r['day'],
                    'date': r['date'],
                    'weekday': r['weekday'],
                    'gio_vao': '\n'.join(vaos),
                    'gio_ra': '\n'.join(ras),
                    'is_absent': is_absent,
                    'missing_checkin': missing_checkin,
                    'missing_checkout': missing_checkout
                })
            
            pdata['records'] = records
            self._persons_data.append(pdata)
    def _normalize_name(self, name: str) -> str:
        """Chuáº©n hÃ³a tÃªn Ä‘á»ƒ so sÃ¡nh"""
        if not name:
            return ""
        name = unicodedata.normalize('NFD', name)
        name = ''.join(c for c in name if unicodedata.category(c) != 'Mn')
        name = re.sub(r'\s+', ' ', name.lower().strip())
        return name

    def _count_work_days(self, person: Dict) -> int:
        """Äáº¿m sá»‘ ngÃ y cÃ³ lÃ m viá»‡c (khÃ´ng vÃ¡ng)"""
        count = 0
        for rec in person.get('records', []):
            if not rec.get('is_absent', True):
                count += 1
        return count

    def _dedupe_persons_by_name(self):
        """
        Náº¿u trÃ¹ng tÃªn, chá»n báº£n ghi cÃ³ sá»‘ ngÃ y Ä‘i lÃ m nhiá»u hÆ¡n.
        Náº¿u báº±ng nhau, chá»n báº£n ghi cÃ³ nhiá»u record hÆ¡n.
        """
        if not self._persons_data:
            return

        picked = {}
        for person in self._persons_data:
            key = self._normalize_name(person.get('name', ''))
            if not key:
                continue

            if key not in picked:
                picked[key] = person
                continue

            current = picked[key]
            work_days = self._count_work_days(person)
            current_work_days = self._count_work_days(current)

            if work_days > current_work_days:
                picked[key] = person
            elif work_days == current_work_days:
                if len(person.get('records', [])) > len(current.get('records', [])):
                    picked[key] = person

        self._persons_data = list(picked.values())
    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def get_persons(self) -> List[Dict]:
        return self._persons_data

    def get_absent_records(self, person: Dict) -> List[Dict]:
        """Láº¥y danh sÃ¡ch ngÃ y váº¯ng hoáº·c thiáº¿u giá» vÃ o/ra"""
        absent = []
        for rec in person['records']:
            if rec['is_absent'] or rec['missing_checkout'] or rec['missing_checkin']:
                issue = 'Váº¯ng máº·t'
                if rec['missing_checkin']:
                    issue = f"Thiáº¿u giá» vÃ o (giá» ra: {rec['gio_ra']})"
                elif rec['missing_checkout']:
                    issue = f"Thiáº¿u giá» ra (giá» vÃ o: {rec['gio_vao']})"
                absent.append({**rec, 'issue': issue})
        return absent


class ExcelToWordExporter:
    """
    Nháº­n dá»¯ liá»‡u tá»« ExcelChamCongExtractor vÃ  táº¡o file Word
    cho tá»«ng ngÆ°á»i, theo format CHI TIáº¾T CHáº¤M CÃ”NG.
    """

    def __init__(
        self,
        portrait_dir: str,
        output_dir: str,
        input_images_dir: str = None,
        face_matcher=None,
        accuracy_mode: bool = False,
        match_distance_threshold: Optional[float] = None,
        log_detail: bool = False
    ):
        self.portrait_dir = portrait_dir
        self.output_dir = output_dir
        self.input_images_dir = input_images_dir
        self.face_matcher = face_matcher
        self.accuracy_mode = accuracy_mode
        self.match_distance_threshold = match_distance_threshold
        self.log_detail = log_detail
        os.makedirs(output_dir, exist_ok=True)
        self._portrait_cache = {}
        self._scan_portraits()

    def _scan_portraits(self):
        if not os.path.exists(self.portrait_dir):
            return
        exts = {'.jpg', '.jpeg', '.png', '.bmp'}
        for item in os.listdir(self.portrait_dir):
            item_path = os.path.join(self.portrait_dir, item)
            if os.path.isdir(item_path):
                for f in os.listdir(item_path):
                    if os.path.splitext(f)[1].lower() in exts:
                        key = self._normalize(item)
                        self._portrait_cache.setdefault(key, []).append(
                            os.path.join(item_path, f))
            else:
                if os.path.splitext(item)[1].lower() in exts:
                    key = self._normalize(os.path.splitext(item)[0])
                    self._portrait_cache.setdefault(key, []).append(item_path)

    def _normalize(self, name: str) -> str:
        import unicodedata
        name = unicodedata.normalize('NFD', name)
        name = ''.join(c for c in name if unicodedata.category(c) != 'Mn')
        return re.sub(r'\s+', ' ', name.lower().strip())

    def _find_portrait(self, name: str) -> Optional[str]:
        key = self._normalize(name)
        if key in self._portrait_cache:
            return self._portrait_cache[key][0]
        for cached_key, paths in self._portrait_cache.items():
            if key in cached_key or cached_key in key:
                return paths[0]
        # Word-by-word match
        words = set(key.split())
        best, best_score = None, 0
        for cached_key, paths in self._portrait_cache.items():
            score = len(words & set(cached_key.split()))
            if score > best_score:
                best_score = score
                best = paths[0]
        return best if best_score >= 2 else None

    def export_person(self, person: Dict, log_callback=None) -> str:
        """Xuáº¥t file Word cho má»™t ngÆ°á»i, tráº£ vá» Ä‘Æ°á»ng dáº«n file"""
        name = person['name']
        month = person['month']
        year = person['year']
        records = person['records']

        doc = Document()

        # === Page margins ===
        section = doc.sections[0]
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)

        # === TiÃªu Ä‘á» ===
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('CHI TIáº¾T CHáº¤M CÃ”NG')
        run.bold = True
        run.font.size = Pt(14)

        # ThÃ´ng tin ngÆ°á»i
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_para.add_run(
            f'TÃªn nhÃ¢n viÃªn: {name.upper()}     '
            f'ThÃ¡ng: {month:02d}/{year}'
        )
        info_run.bold = True
        info_run.font.size = Pt(11)

        doc.add_paragraph()

        # === Báº£ng cháº¥m cÃ´ng ===
        headers = ['STT', 'NgÃ y', 'Thá»©', 'Giá» vÃ o', 'Giá» ra', 'Ghi chÃº', 'áº¢nh Camera']
        col_widths = [Cm(1.0), Cm(2.7), Cm(1.5), Cm(2.2), Cm(2.2), Cm(3.2), Cm(3.5)]

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        hdr_cells = table.rows[0].cells
        for i, (h, w) in enumerate(zip(headers, col_widths)):
            hdr_cells[i].width = w
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(10)

        # Data rows
        absent_days = []
        for stt, rec in enumerate(records, 1):
            row = table.add_row()
            cells = row.cells

            # Highlight absent / missing rows
            is_issue = rec['is_absent'] or rec['missing_checkout'] or rec['missing_checkin']
            if is_issue:
                absent_days.append(rec)

            data = [
                str(stt),
                rec['date'],
                rec['weekday'],
                rec['gio_vao'] or '',
                rec['gio_ra'] or '',
            ]

            note = ''
            if rec['is_absent']:
                note = 'Váº¯ng máº·t'
            elif rec['missing_checkin']:
                note = 'Thiáº¿u giá» vÃ o'
            elif rec['missing_checkout']:
                note = 'Thiáº¿u giá» ra'

            data.append(note)

            for i, (val, w) in enumerate(zip(data, col_widths[:6])):
                cells[i].width = w
                p = cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 5 else WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(val)
                run.font.size = Pt(9)
                if is_issue and note:
                    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)  # Äá»
            
            # Xá»­ lÃ½ áº£nh camera cho cá»™t cuá»‘i cÃ¹ng
            cells[6].width = col_widths[6]
            if is_issue and self.input_images_dir and self.face_matcher:
                day_str = rec['date'].split('/')[0].zfill(2)
                day_folder = os.path.join(self.input_images_dir, day_str)
                matched_img = None
                
                if os.path.exists(day_folder):
                    # TÃ¬m táº¥t cáº£ file áº£nh trong thÆ° má»¥c ngÃ y nÃ y
                    camera_images = []
                    for root, _, files in os.walk(day_folder):
                        for f in files:
                            if os.path.splitext(f)[1].lower() in ['.jpg', '.jpeg', '.png', '.bmp']:
                                camera_images.append(os.path.join(root, f))
                    
                    if camera_images:
                        if log_callback:
                            log_callback(f"    ðŸ” {name} (ngÃ y {day_str}): So sÃ¡nh {len(camera_images)} áº£nh...", "default")
                        try:
                            matched_img = self.face_matcher.match_face_in_images(
                                name,
                                camera_images,
                                distance_threshold=self.match_distance_threshold,
                                fast_mode=not self.accuracy_mode,
                                log_detail=self.log_detail
                            )
                        except Exception as e:
                            if log_callback:
                                log_callback(f"    âŒ Lá»—i OpenCV/DeepFace khi so sÃ¡nh {name}: {e}", "warning")
                    else:
                        if log_callback:
                            log_callback(f"    âš ï¸ {name} (ngÃ y {day_str}): ThÆ° má»¥c áº£nh rá»—ng", "warning")
                else:
                    if log_callback:
                        log_callback(f"    âš ï¸ {name} (ngÃ y {day_str}): KhÃ´ng cÃ³ thÆ° má»¥c áº£nh", "warning")
                
                # ChÃ¨n áº£nh vÃ o Ã´
                p = cells[6].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if matched_img and os.path.exists(matched_img):
                    try:
                        run = p.add_run()
                        run.add_picture(matched_img, width=Cm(3))
                        if log_callback:
                            log_callback(f"    âœ… {name}: ÄÃ£ chÃ¨n áº£nh {os.path.basename(matched_img)}", "success")
                    except Exception as e:
                        p.add_run(f"[Lá»—i áº£nh]").font.size = Pt(8)
                else:
                    if os.path.exists(day_folder) and camera_images:
                        p.add_run("KhÃ´ng khá»›p").font.size = Pt(8)
                    elif not os.path.exists(day_folder):
                        p.add_run("KhÃ´ng cÃ³ dl").font.size = Pt(8)
            else:
                p = cells[6].paragraphs[0]
                run = p.add_run()
                run.font.size = Pt(9)

        # === TÃ³m táº¯t váº¯ng ===
        if absent_days:
            doc.add_paragraph()
            summary_para = doc.add_paragraph()
            summary_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            sr = summary_para.add_run(f'Tá»•ng ngÃ y váº¯ng/thiáº¿u dá»¯ liá»‡u: {len(absent_days)} ngÃ y')
            sr.bold = True
            sr.font.size = Pt(10)

        # === Save ===
        safe_name = re.sub(r'[<>:"/\\|?*]', '_', name.strip().upper())
        filename = f'{safe_name}_{month:02d}_{year}.docx'
        output_path = os.path.join(self.output_dir, filename)
        doc.save(output_path)

        if log_callback:
            log_callback(f'âœ… ÄÃ£ xuáº¥t: {filename} ({len(records)} ngÃ y, {len(absent_days)} váº¯ng/thiáº¿u)')

        return output_path

    def export_all(self, persons: List[Dict], log_callback=None) -> List[str]:
        """Xuáº¥t file Word cho táº¥t cáº£ má»i ngÆ°á»i"""
        results = []
        for person in persons:
            try:
                path = self.export_person(person, log_callback=log_callback)
                results.append(path)
            except Exception as e:
                if log_callback:
                    log_callback(f'âŒ Lá»—i xuáº¥t {person["name"]}: {e}', 'error')
        return results


# â”€â”€â”€ CLI test â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
if __name__ == '__main__':
    import sys
    sys.stdout.reconfigure(encoding='utf-8')

    excel_path = r'd:\Projects\phan mem quet mat\000001_02BÃ¡o cÃ¡o 28.02.2026.xls'
    portrait_dir = r'd:\Projects\phan mem quet mat\áº¢nh BV'
    output_dir = r'd:\Projects\phan mem quet mat\results\excel_export'

    print('ðŸ“‚ Äang Ä‘á»c file Excel...')
    extractor = ExcelChamCongExtractor(excel_path)
    persons = extractor.get_persons()
    print(f'âœ… TÃ¬m tháº¥y {len(persons)} ngÆ°á»i')
    for p in persons[:5]:
        absent = extractor.get_absent_records(p)
        print(f'  - {p["name"]}: {len(p["records"])} ngÃ y, {len(absent)} váº¯ng/thiáº¿u')

    print('\nðŸ“ Äang xuáº¥t file Word...')
    exporter = ExcelToWordExporter(portrait_dir, output_dir)
    files = exporter.export_all(persons, log_callback=print)
    print(f'\nðŸŽ‰ ÄÃ£ xuáº¥t {len(files)} file Word vÃ o: {output_dir}')


