# -*- coding: utf-8 -*-
"""
Module xử lý chấm công - đọc file Word và phát hiện ngày thiếu check-in/out
"""

import os
import re
from datetime import datetime
from docx import Document
from typing import List, Dict, Optional, Tuple


class AttendanceProcessor:
    """Xử lý file Word chấm công và phát hiện ngày vắng/thiếu dữ liệu"""
    
    def __init__(self, chamcong_dir: str):
        self.chamcong_dir = chamcong_dir
        self.attendance_data = {}  # {person_name: [attendance_records]}
        self.missing_records = []  # Danh sách thiếu dữ liệu
        
    def scan_all_files(self) -> Dict[str, List[Dict]]:
        """Quét tất cả file Word trong thư mục chấm công"""
        if not os.path.exists(self.chamcong_dir):
            print(f"Thư mục không tồn tại: {self.chamcong_dir}")
            return {}
        
        for filename in os.listdir(self.chamcong_dir):
            if filename.endswith('.docx') and not filename.startswith('~$'):
                # Loại bỏ file giải trình tổng hợp
                if 'GIẢI TRÌNH' in filename.upper() or 'GIAI TRINH' in filename.upper():
                    continue
                    
                filepath = os.path.join(self.chamcong_dir, filename)
                person_name = os.path.splitext(filename)[0]
                
                try:
                    records = self._parse_attendance_file(filepath)
                    self.attendance_data[person_name] = records
                except Exception as e:
                    print(f"Lỗi đọc file {filename}: {e}")
        
        return self.attendance_data
    
    def _parse_attendance_file(self, filepath: str) -> List[Dict]:
        """Parse file Word chấm công, trả về danh sách các bản ghi"""
        doc = Document(filepath)
        records = []
        
        if not doc.tables:
            return records
        
        table = doc.tables[0]
        
        # Tìm dòng bắt đầu dữ liệu (sau header)
        data_start_row = 8  # Thường dữ liệu bắt đầu từ row 8
        
        for row_idx, row in enumerate(table.rows[data_start_row:], start=data_start_row):
            cells = [cell.text.strip() for cell in row.cells]
            
            if len(cells) < 16:
                continue
            
            # Cột: Ngày, Thứ, Vào1, Ra1, Vào2, Ra2, Vào3, Ra3, Trễ, Sớm, Về trễ, Giờ, Công, T.Ca1, T.Ca2, Ký hiệu
            date_str = cells[0]
            weekday = cells[1]
            vao1, ra1 = cells[2], cells[3]
            vao2, ra2 = cells[4], cells[5]
            vao3, ra3 = cells[6], cells[7]
            ky_hieu = cells[15] if len(cells) > 15 else ''
            
            # Bỏ qua dòng không phải ngày hợp lệ
            if not self._is_valid_date(date_str):
                continue
            
            # Kiểm tra trạng thái (nhưng vẫn kiểm tra vấn đề)
            is_off_day = ky_hieu.upper() in ['V', 'P', 'NL']  # Nghỉ, Phép, Nghỉ lễ
            
            record = {
                'date': date_str,
                'weekday': weekday,
                'check_in_1': vao1,
                'check_out_1': ra1,
                'check_in_2': vao2,
                'check_out_2': ra2,
                'check_in_3': vao3,
                'check_out_3': ra3,
                'symbol': ky_hieu,
                'is_off_day': is_off_day,
                'has_issue': False,
                'issue_type': None
            }
            
            # Phát hiện vấn đề - luôn kiểm tra cả ngày nghỉ nếu có text không hợp lệ
            record = self._detect_issues(record)
            
            records.append(record)
        
        return records
    
    def _is_valid_date(self, date_str: str) -> bool:
        """Kiểm tra chuỗi có phải ngày hợp lệ không (dd/mm/yyyy)"""
        pattern = r'^\d{1,2}/\d{1,2}/\d{4}$'
        return bool(re.match(pattern, date_str))
    
    def _detect_issues(self, record: Dict) -> Dict:
        """Phát hiện vấn đề trong bản ghi chấm công"""
        vao1, ra1 = record['check_in_1'], record['check_out_1']
        vao2, ra2 = record['check_in_2'], record['check_out_2']
        vao3, ra3 = record['check_in_3'], record['check_out_3']
        
        # Kiểm tra có giờ hợp lệ không
        has_valid_checkin = self._is_valid_time(vao1) or self._is_valid_time(vao2) or self._is_valid_time(vao3)
        has_valid_checkout = self._is_valid_time(ra1) or self._is_valid_time(ra2) or self._is_valid_time(ra3)
        
        # Kiểm tra có text không hợp lệ (không phải số giờ)
        all_time_fields = [vao1, ra1, vao2, ra2, vao3, ra3]
        has_invalid_text = any(
            field and field.strip() and not self._is_valid_time(field) 
            for field in all_time_fields
        )
        
        # Phát hiện các loại vấn đề
        if has_invalid_text:
            # Có text không hợp lệ (như "Nghỉ", "Chuyển", "ca", v.v.)
            record['has_issue'] = True
            record['issue_type'] = 'invalid_text'
            record['invalid_values'] = [f for f in all_time_fields if f and f.strip() and not self._is_valid_time(f)]
        elif not has_valid_checkin and not has_valid_checkout:
            record['has_issue'] = True
            record['issue_type'] = 'missing_both'  # Thiếu cả vào và ra
        elif not has_valid_checkin:
            record['has_issue'] = True
            record['issue_type'] = 'missing_checkin'  # Thiếu giờ vào
        elif not has_valid_checkout:
            record['has_issue'] = True
            record['issue_type'] = 'missing_checkout'  # Thiếu giờ ra
        
        return record
    
    def _is_valid_time(self, time_str: str) -> bool:
        """Kiểm tra chuỗi có phải giờ hợp lệ không (HH:MM)"""
        if not time_str:
            return False
        pattern = r'^\d{1,2}:\d{2}$'
        return bool(re.match(pattern, time_str.strip()))
    
    def get_missing_records(self) -> List[Dict]:
        """Lấy danh sách tất cả các bản ghi thiếu dữ liệu"""
        missing = []
        
        for person_name, records in self.attendance_data.items():
            for record in records:
                if record['has_issue']:
                    issue_desc = {
                        'missing_both': 'Thiếu giờ vào và ra',
                        'missing_checkin': 'Thiếu giờ vào',
                        'missing_checkout': 'Thiếu giờ ra',
                        'invalid_text': 'Dữ liệu không hợp lệ: ' + ', '.join(record.get('invalid_values', []))
                    }.get(record['issue_type'], 'Thiếu dữ liệu')
                    
                    missing.append({
                        'person_name': person_name,
                        'date': record['date'],
                        'weekday': record['weekday'],
                        'issue_type': record['issue_type'],
                        'issue_description': issue_desc,
                        'explanation': 'Nhân viên có trực, bổ sung'  # Mặc định
                    })
        
        # Sắp xếp theo tên, rồi theo ngày
        missing.sort(key=lambda x: (x['person_name'], x['date']))
        self.missing_records = missing
        return missing
    
    def get_summary(self) -> Dict:
        """Tóm tắt kết quả quét"""
        total_persons = len(self.attendance_data)
        total_records = sum(len(r) for r in self.attendance_data.values())
        total_missing = len(self.get_missing_records())
        
        # Đếm theo loại issue
        issue_counts = {}
        for record in self.missing_records:
            issue_type = record['issue_type']
            issue_counts[issue_type] = issue_counts.get(issue_type, 0) + 1
        
        return {
            'total_persons': total_persons,
            'total_records': total_records,
            'total_missing': total_missing,
            'issue_breakdown': issue_counts,
            'persons_with_issues': len(set(r['person_name'] for r in self.missing_records))
        }


# Test
if __name__ == '__main__':
    import sys
    sys.stdout.reconfigure(encoding='utf-8')
    
    chamcong_dir = r'd:\Projects\phan mem quet mat\chamcong'
    processor = AttendanceProcessor(chamcong_dir)
    
    print("Đang quét file chấm công...")
    processor.scan_all_files()
    
    print("\n=== Tóm tắt ===")
    summary = processor.get_summary()
    print(f"Tổng số người: {summary['total_persons']}")
    print(f"Tổng số bản ghi: {summary['total_records']}")
    print(f"Tổng số thiếu: {summary['total_missing']}")
    print(f"Số người có vấn đề: {summary['persons_with_issues']}")
    
    print("\n=== Danh sách thiếu (10 đầu tiên) ===")
    missing = processor.get_missing_records()
    for i, rec in enumerate(missing[:10], 1):
        print(f"{i}. {rec['person_name']} - {rec['date']} ({rec['weekday']}): {rec['issue_description']}")
