# -*- coding: utf-8 -*-
"""
Flask Web Server cho pháº§n má»m nháº­n diá»‡n khuÃ´n máº·t
PhiÃªn báº£n Ä‘Æ¡n giáº£n Ä‘á»ƒ test UI
"""

import os
import sys
import time
import json
import queue
import threading
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed

# ThÃªm thÆ° má»¥c gá»‘c vÃ o path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from flask import Flask, render_template, request, jsonify, send_file, send_from_directory, Response
from werkzeug.utils import secure_filename

# Cáº¥u hÃ¬nh - PhÃ¡t hiá»‡n Ä‘Ãºng thÆ° má»¥c khi cháº¡y tá»« EXE
def get_base_dir():
    """Láº¥y thÆ° má»¥c gá»‘c (chá»©a data: input_images, database, ...) - há»— trá»£ cáº£ khi cháº¡y tá»« source vÃ  tá»« EXE"""
    if getattr(sys, 'frozen', False):
        # Cháº¡y tá»« EXE (PyInstaller) - data náº±m cáº¡nh file exe
        return os.path.dirname(sys.executable)
    else:
        # Cháº¡y tá»« source
        return os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

def get_resource_dir():
    """Láº¥y thÆ° má»¥c resources (templates, static) - há»— trá»£ cáº£ khi cháº¡y tá»« EXE"""
    if getattr(sys, 'frozen', False):
        # Cháº¡y tá»« EXE - resources Ä‘Æ°á»£c PyInstaller giáº£i nÃ©n vÃ o _MEIPASS
        return sys._MEIPASS
    else:
        # Cháº¡y tá»« source
        return os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

BASE_DIR = get_base_dir()
RESOURCE_DIR = get_resource_dir()
INPUT_IMAGES_DIR = os.path.join(BASE_DIR, "input_images")
DATABASE_DIR = os.path.join(BASE_DIR, "database")
RESULTS_DIR = os.path.join(BASE_DIR, "results")

def _normalize_folder_name(name: str) -> str:
    import unicodedata
    name = unicodedata.normalize('NFD', name)
    name = ''.join(c for c in name if unicodedata.category(c) != 'Mn')
    return ' '.join(name.lower().strip().split())

def _count_images_in_dir(path: str) -> int:
    exts = {'.jpg', '.jpeg', '.png', '.bmp'}
    count = 0
    for root, _, files in os.walk(path):
        for f in files:
            if os.path.splitext(f)[1].lower() in exts:
                count += 1
    return count

def resolve_portrait_dir(base_dir: str) -> str:
    """Find portrait folder even if accents/encoding differ."""
    candidates = ["Ảnh BV", "Anh BV", "ANH BV", "anh bv", "ẢnhBV"]
    for c in candidates:
        p = os.path.join(base_dir, c)
        if os.path.exists(p):
            return p
    try:
        target = _normalize_folder_name("Ảnh BV")
        for name in os.listdir(base_dir):
            if _normalize_folder_name(name) == target:
                return os.path.join(base_dir, name)
    except Exception:
        pass
    return os.path.join(base_dir, "Ảnh BV")

def resolve_portrait_dir_by_scan(base_dir: str) -> str:
    """Fallback: scan subfolders and pick one that looks like portrait dir."""
    best_dir = None
    best_count = 0
    keywords = ("anh", "chan dung", "portrait", "bv")
    try:
        for name in os.listdir(base_dir):
            path = os.path.join(base_dir, name)
            if not os.path.isdir(path):
                continue
            norm = _normalize_folder_name(name)
            if not any(k in norm for k in keywords):
                continue
            count = _count_images_in_dir(path)
            if count > best_count:
                best_count = count
                best_dir = path
    except Exception:
        pass
    return best_dir or resolve_portrait_dir(base_dir)
SUPPORTED_IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.bmp', '.gif'}

FLASK_HOST = "0.0.0.0"
FLASK_PORT = 5000
FLASK_DEBUG = True
MAX_WORKERS = 4

# Thu muc du lieu duoc tao tu dong khi clone moi
CHAMCONG_DIR = os.path.join(BASE_DIR, "chamcong")
PORTRAIT_DIR = resolve_portrait_dir(BASE_DIR)
NGAY_RONG_DIR = os.path.join(BASE_DIR, "ngay_rong")

# Tao thu muc neu chua ton tai
for directory in [INPUT_IMAGES_DIR, DATABASE_DIR, RESULTS_DIR, CHAMCONG_DIR, PORTRAIT_DIR, NGAY_RONG_DIR]:
    os.makedirs(directory, exist_ok=True)

# Flask app - dÃ¹ng RESOURCE_DIR cho templates/static (Ä‘Ãºng cáº£ khi cháº¡y tá»« EXE)
app = Flask(__name__, 
            template_folder=os.path.join(RESOURCE_DIR, 'templates'),
            static_folder=os.path.join(RESOURCE_DIR, 'static'))

app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max upload

# ==================== TASK MANAGER ====================

class ProcessingTask:
    def __init__(self, task_id):
        self.task_id = task_id
        self.status = 'pending'
        self.progress = 0
        self.total = 0
        self.current_file = ''
        self.results = []
        self.errors = []
        self.start_time = None
        self.end_time = None
        self.output_file = None
    
    def to_dict(self):
        return {
            'task_id': self.task_id,
            'status': self.status,
            'progress': self.progress,
            'total': self.total,
            'current_file': self.current_file,
            'results_count': len(self.results),
            'errors_count': len(self.errors),
            'start_time': self.start_time.isoformat() if self.start_time else None,
            'end_time': self.end_time.isoformat() if self.end_time else None,
            'output_file': self.output_file
        }

# Global state
tasks = {}
database = {}

# ==================== LOG STREAMING ====================
# Global log queue for SSE streaming
log_queues = []
log_lock = threading.Lock()

# Setup file logging
LOG_FILE = os.path.join(BASE_DIR, 'app.log')

def setup_file_logging():
    """Configure logging to file"""
    import logging
    logging.basicConfig(
        filename=LOG_FILE,
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        encoding='utf-8'
    )
    # Log startup
    logging.info("="*50)
    logging.info("APP STARTED")
    logging.info(f"Frozen: {getattr(sys, 'frozen', False)}")
    logging.info(f"Base Dir (data): {BASE_DIR}")
    logging.info(f"Resource Dir (templates/static): {RESOURCE_DIR}")
    logging.info(f"Input Images Dir: {INPUT_IMAGES_DIR}")
    logging.info(f"Template folder: {app.template_folder}")
    logging.info(f"Static folder: {app.static_folder}")
    logging.info(f"Template folder exists: {os.path.exists(app.template_folder)}")
    logging.info(f"Static folder exists: {os.path.exists(app.static_folder)}")
    logging.info("="*50)

# Initialize logging
setup_file_logging()

def send_log(message, log_type='default'):
    """Send log message to all connected SSE clients and write to file"""
    import logging
    
    # Write to file
    if log_type == 'error':
        logging.error(message)
    elif log_type == 'warning':
        logging.warning(message)
    elif log_type == 'success':
        logging.info(f"[SUCCESS] {message}")
    else:
        logging.info(message)
        
    # Also print to console
    print(message)
    
    # Send to SSE clients
    with log_lock:
        for q in log_queues:
            try:
                msg_data = json.dumps({
                    'message': message,
                    'type': log_type,
                    'time': datetime.now().strftime('%H:%M:%S')
                })
                q.put(msg_data)
            except Exception:
                pass


def get_image_files(folder_path):
    """Láº¥y danh sÃ¡ch file áº£nh trong thÆ° má»¥c"""
    image_files = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            ext = os.path.splitext(file)[1].lower()
            if ext in SUPPORTED_IMAGE_EXTENSIONS:
                image_files.append(os.path.join(root, file))
    return image_files

def process_image(image_path):
    """Xá»­ lÃ½ má»™t áº£nh (demo version - tráº£ vá» dá»¯ liá»‡u máº«u)"""
    result = {
        'image_path': image_path,
        'filename': os.path.basename(image_path),
        'datetime': None,
        'location': None,
        'faces': [],
        'matched_person': None,
        'branch': None,
        'person_name': None,
        'confidence': None,
        'error': None
    }
    
    # Simulate processing time
    time.sleep(0.5)
    
    # Demo data
    result['datetime'] = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
    result['location'] = 'Demo Location'
    result['faces'] = [{'location': (0, 100, 100, 0)}]
    
    # TÃ¬m trong database
    if database:
        for person_id, data in database.items():
            result['matched_person'] = person_id
            result['branch'] = data.get('branch', 'Unknown')
            result['person_name'] = data.get('name', 'Unknown')
            result['confidence'] = 85.5
            break
    
    return result

def run_processing(task_id, files):
    """Xá»­ lÃ½ trong background thread"""
    task = tasks[task_id]
    task.status = 'running'
    task.start_time = datetime.now()
    task.total = len(files)
    
    try:
        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
            futures = {executor.submit(process_image, f): f for f in files}
            
            for future in as_completed(futures):
                file_path = futures[future]
                task.current_file = os.path.basename(file_path)
                
                try:
                    result = future.result(timeout=60)
                    task.results.append(result)
                except Exception as e:
                    task.errors.append({'file': file_path, 'error': str(e)})
                
                task.progress += 1
        
        # Export to Excel
        task.output_file = export_results(task)
        task.status = 'completed'
        
    except Exception as e:
        task.status = 'failed'
        task.errors.append(str(e))
    
    task.end_time = datetime.now()

def export_results(task):
    """Xuáº¥t káº¿t quáº£ ra file Excel hoáº·c CSV"""
    try:
        import pandas as pd
        
        rows = []
        for i, result in enumerate(task.results, 1):
            rows.append({
                'STT': i,
                'TÃªn File': result['filename'],
                'NgÃ y Giá»': result['datetime'] or '',
                'Äá»‹a Äiá»ƒm': result['location'] or '',
                'Chi NhÃ¡nh': result['branch'] or '',
                'TÃªn NgÆ°á»i': result['person_name'] or 'KhÃ´ng xÃ¡c Ä‘á»‹nh',
                'Äá»™ Tin Cáº­y (%)': result['confidence'] or 0,
                'Sá»‘ KhuÃ´n Máº·t': len(result['faces']),
                'Lá»—i': result['error'] or ''
            })
        
        df = pd.DataFrame(rows)
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f'result_{timestamp}.xlsx'
        output_path = os.path.join(RESULTS_DIR, output_filename)
        
        df.to_excel(output_path, index=False, engine='openpyxl')
        return output_path
        
    except ImportError:
        # Fallback to CSV if pandas not available
        import csv
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f'result_{timestamp}.csv'
        output_path = os.path.join(RESULTS_DIR, output_filename)
        
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['STT', 'TÃªn File', 'NgÃ y Giá»', 'Äá»‹a Äiá»ƒm', 'Chi NhÃ¡nh', 'TÃªn NgÆ°á»i', 'Äá»™ Tin Cáº­y (%)', 'Sá»‘ KhuÃ´n Máº·t', 'Lá»—i'])
            
            for i, result in enumerate(task.results, 1):
                writer.writerow([
                    i,
                    result['filename'],
                    result['datetime'] or '',
                    result['location'] or '',
                    result['branch'] or '',
                    result['person_name'] or 'KhÃ´ng xÃ¡c Ä‘á»‹nh',
                    result['confidence'] or 0,
                    len(result['faces']),
                    result['error'] or ''
                ])
        
        return output_path
    except Exception as e:
        print(f"Lá»—i xuáº¥t file: {e}")
        return None

def scan_database():
    """QuÃ©t database áº£nh chÃ¢n dung"""
    global database
    database = {}
    
    if not os.path.exists(DATABASE_DIR):
        return
    
    for branch in os.listdir(DATABASE_DIR):
        branch_path = os.path.join(DATABASE_DIR, branch)
        if not os.path.isdir(branch_path) or branch.startswith('.'):
            continue
        
        for person_name in os.listdir(branch_path):
            person_path = os.path.join(branch_path, person_name)
            if not os.path.isdir(person_path):
                continue
            
            # TÃ¬m áº£nh
            image_files = []
            for file in os.listdir(person_path):
                ext = os.path.splitext(file)[1].lower()
                if ext in SUPPORTED_IMAGE_EXTENSIONS:
                    image_files.append(os.path.join(person_path, file))
            
            if image_files:
                person_id = f"{branch}/{person_name}"
                database[person_id] = {
                    'encoding': None,  # Placeholder
                    'branch': branch,
                    'name': person_name,
                    'image_path': image_files[0]
                }

# ==================== PAGES ====================

@app.route('/')
def index():
    return render_template('index.html')

# ==================== SSE LOG STREAM ====================

@app.route('/api/log-stream')
def log_stream():
    """SSE endpoint for real-time log streaming"""
    def generate():
        q = queue.Queue()
        with log_lock:
            log_queues.append(q)
        try:
            while True:
                try:
                    # Wait for log message with timeout
                    data = q.get(timeout=30)
                    yield f"data: {data}\n\n"
                except queue.Empty:
                    # Send heartbeat to keep connection alive
                    yield f"data: {json.dumps({'message': '', 'type': 'heartbeat'})}\n\n"
        except GeneratorExit:
            pass
        finally:
            with log_lock:
                if q in log_queues:
                    log_queues.remove(q)
    
    return Response(generate(), mimetype='text/event-stream', 
                    headers={'Cache-Control': 'no-cache', 'Connection': 'keep-alive'})

# ==================== API: SCAN ====================

@app.route('/api/scan/start', methods=['POST'])
def start_scan():
    data = request.json or {}
    date_folder = data.get('date', None)  # CÃ³ thá»ƒ chá»‰ Ä‘á»‹nh ngÃ y cá»¥ thá»ƒ
    folder_path = data.get('folder_path', INPUT_IMAGES_DIR)
    
    # Náº¿u cÃ³ chá»‰ Ä‘á»‹nh ngÃ y, quÃ©t trong thÆ° má»¥c ngÃ y Ä‘Ã³
    if date_folder:
        folder_path = os.path.join(INPUT_IMAGES_DIR, date_folder)
    
    if not os.path.exists(folder_path):
        return jsonify({'error': f'ThÆ° má»¥c khÃ´ng tá»“n táº¡i: {folder_path}'}), 400
    
    files = get_image_files(folder_path)
    if not files:
        return jsonify({'error': 'KhÃ´ng tÃ¬m tháº¥y file áº£nh nÃ o'}), 400
    
    task_id = f"task_{int(time.time() * 1000)}"
    task = ProcessingTask(task_id)
    tasks[task_id] = task
    
    thread = threading.Thread(target=run_processing, args=(task_id, files))
    thread.daemon = True
    thread.start()
    
    return jsonify({
        'success': True,
        'task_id': task_id,
        'message': f'ÄÃ£ báº¯t Ä‘áº§u quÃ©t {len(files)} áº£nh',
        'folder': folder_path,
        'total_files': len(files)
    })

@app.route('/api/scan/dates')
def list_date_folders():
    """Liá»‡t kÃª cÃ¡c thÆ° má»¥c ngÃ y cÃ³ sáºµn trong input_images"""
    date_folders = []
    if os.path.exists(INPUT_IMAGES_DIR):
        for item in os.listdir(INPUT_IMAGES_DIR):
            item_path = os.path.join(INPUT_IMAGES_DIR, item)
            if os.path.isdir(item_path):
                # Äáº¿m sá»‘ áº£nh trong thÆ° má»¥c
                image_count = len(get_image_files(item_path))
                if image_count > 0:
                    date_folders.append({
                        'name': item,
                        'path': item_path,
                        'image_count': image_count
                    })
    
    # Sáº¯p xáº¿p theo tÃªn (ngÃ y)
    date_folders.sort(key=lambda x: x['name'])
    
    return jsonify({
        'success': True,
        'folders': date_folders,
        'total': len(date_folders)
    })

@app.route('/api/scan/all-dates', methods=['POST'])
def scan_all_dates():
    """QuÃ©t táº¥t cáº£ cÃ¡c thÆ° má»¥c ngÃ y cÃ³ áº£nh trong input_images"""
    date_folders = []
    total_images = 0
    
    if not os.path.exists(INPUT_IMAGES_DIR):
        return jsonify({'error': 'ThÆ° má»¥c input_images khÃ´ng tá»“n táº¡i'}), 400
    
    # TÃ¬m táº¥t cáº£ thÆ° má»¥c con cÃ³ áº£nh
    for item in os.listdir(INPUT_IMAGES_DIR):
        item_path = os.path.join(INPUT_IMAGES_DIR, item)
        if os.path.isdir(item_path):
            images = get_image_files(item_path)
            if images:
                date_folders.append({
                    'name': item,
                    'path': item_path,
                    'images': images
                })
                total_images += len(images)
    
    if not date_folders:
        return jsonify({'error': 'KhÃ´ng tÃ¬m tháº¥y thÆ° má»¥c ngÃ y nÃ o cÃ³ áº£nh'}), 400
    
    # Sáº¯p xáº¿p theo tÃªn ngÃ y
    date_folders.sort(key=lambda x: x['name'])
    
    # Táº¡o task vÃ  báº¯t Ä‘áº§u quÃ©t tá»«ng thÆ° má»¥c
    task_id = f"task_{int(time.time() * 1000)}"
    task = ProcessingTask(task_id)
    task.total = total_images
    tasks[task_id] = task
    
    def process_all_dates():
        task.status = 'running'
        task.start_time = datetime.now()
        
        for folder_info in date_folders:
            folder_name = folder_info['name']
            images = folder_info['images']
            
            for image_path in images:
                task.current_file = f"[{folder_name}] {os.path.basename(image_path)}"
                
                try:
                    result = process_image(image_path)
                    result['date_folder'] = folder_name  # ThÃªm thÃ´ng tin thÆ° má»¥c ngÃ y
                    task.results.append(result)
                except Exception as e:
                    task.errors.append({'file': image_path, 'error': str(e)})
                
                task.progress += 1
        
        # Export káº¿t quáº£
        task.output_file = export_results(task)
        task.status = 'completed'
        task.end_time = datetime.now()
    
    thread = threading.Thread(target=process_all_dates, daemon=True)
    thread.start()
    
    return jsonify({
        'success': True,
        'task_id': task_id,
        'message': f'Äang quÃ©t {len(date_folders)} thÆ° má»¥c, tá»•ng {total_images} áº£nh',
        'folders': [{'name': f['name'], 'count': len(f['images'])} for f in date_folders],
        'total_images': total_images
    })

@app.route('/api/scan/status/<task_id>')
def get_scan_status(task_id):
    task = tasks.get(task_id)
    if task:
        return jsonify(task.to_dict())
    return jsonify({'error': 'Task khÃ´ng tá»“n táº¡i'}), 404

@app.route('/api/scan/results/<task_id>')
def get_scan_results(task_id):
    task = tasks.get(task_id)
    if task:
        return jsonify({
            **task.to_dict(),
            'results': task.results,
            'errors': task.errors
        })
    return jsonify({'error': 'Task khÃ´ng tá»“n táº¡i'}), 404

@app.route('/api/scan/tasks')
def get_all_tasks():
    return jsonify({'tasks': [t.to_dict() for t in tasks.values()]})

# ==================== API: DATABASE ====================

@app.route('/api/database/stats')
def get_database_stats():
    stats = {
        'total_persons': len(database),
        'total_branches': len(set(d['branch'] for d in database.values())),
        'branches': {}
    }
    
    for person_id, data in database.items():
        branch = data.get('branch', 'Unknown')
        if branch not in stats['branches']:
            stats['branches'][branch] = 0
        stats['branches'][branch] += 1
    
    return jsonify(stats)

@app.route('/api/database/scan', methods=['POST'])
def rescan_database():
    scan_database()
    stats = get_database_stats().get_json()
    return jsonify({
        'success': True,
        'message': 'ÄÃ£ quÃ©t database',
        'stats': stats
    })

@app.route('/api/database/branches')
def get_branches():
    branches = []
    if os.path.exists(DATABASE_DIR):
        for branch in os.listdir(DATABASE_DIR):
            branch_path = os.path.join(DATABASE_DIR, branch)
            if os.path.isdir(branch_path) and not branch.startswith('.'):
                branches.append(branch)
    return jsonify({'branches': branches})

@app.route('/api/database/branches', methods=['POST'])
def add_branch():
    data = request.json or {}
    branch_name = data.get('name', '').strip()
    
    if not branch_name:
        return jsonify({'error': 'TÃªn chi nhÃ¡nh khÃ´ng há»£p lá»‡'}), 400
    
    branch_path = os.path.join(DATABASE_DIR, branch_name)
    os.makedirs(branch_path, exist_ok=True)
    
    return jsonify({
        'success': True,
        'message': f'ÄÃ£ thÃªm chi nhÃ¡nh: {branch_name}'
    })

@app.route('/api/database/persons/<branch>')
def get_persons(branch):
    persons = []
    branch_path = os.path.join(DATABASE_DIR, branch)
    
    if os.path.exists(branch_path):
        for person in os.listdir(branch_path):
            person_path = os.path.join(branch_path, person)
            if os.path.isdir(person_path):
                persons.append({
                    'name': person,
                    'branch': branch,
                    'person_id': f"{branch}/{person}"
                })
    
    return jsonify({'persons': persons})

# ==================== API: FILES ====================

@app.route('/api/files/input')
def list_input_files():
    files = []
    if os.path.exists(INPUT_IMAGES_DIR):
        for file in os.listdir(INPUT_IMAGES_DIR):
            file_path = os.path.join(INPUT_IMAGES_DIR, file)
            if os.path.isfile(file_path):
                ext = os.path.splitext(file)[1].lower()
                if ext in SUPPORTED_IMAGE_EXTENSIONS:
                    files.append({
                        'name': file,
                        'size': os.path.getsize(file_path),
                        'path': file_path
                    })
    
    return jsonify({
        'folder': INPUT_IMAGES_DIR,
        'files': files,
        'count': len(files)
    })

@app.route('/api/files/results')
def list_result_files():
    files = []
    if os.path.exists(RESULTS_DIR):
        for file in os.listdir(RESULTS_DIR):
            if file.endswith('.xlsx') or file.endswith('.csv'):
                file_path = os.path.join(RESULTS_DIR, file)
                files.append({
                    'name': file,
                    'size': os.path.getsize(file_path),
                    'path': file_path
                })
    
    files.sort(key=lambda x: x['name'], reverse=True)
    
    return jsonify({
        'folder': RESULTS_DIR,
        'files': files,
        'count': len(files)
    })

@app.route('/api/files/download/<filename>')
def download_file(filename):
    file_path = os.path.join(RESULTS_DIR, secure_filename(filename))
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    return jsonify({'error': 'File khÃ´ng tá»“n táº¡i'}), 404

@app.route('/api/files/upload', methods=['POST'])
def upload_files():
    if 'files' not in request.files:
        return jsonify({'error': 'KhÃ´ng cÃ³ file Ä‘Æ°á»£c upload'}), 400
    
    files = request.files.getlist('files')
    uploaded = []
    
    for file in files:
        if file.filename:
            filename = secure_filename(file.filename)
            ext = os.path.splitext(filename)[1].lower()
            
            if ext in SUPPORTED_IMAGE_EXTENSIONS:
                file_path = os.path.join(INPUT_IMAGES_DIR, filename)
                file.save(file_path)
                uploaded.append(filename)
    
    return jsonify({
        'success': True,
        'uploaded': uploaded,
        'count': len(uploaded)
    })

# ==================== API: CONFIG ====================

@app.route('/api/config')
def get_config():
    return jsonify({
        'input_dir': INPUT_IMAGES_DIR,
        'database_dir': DATABASE_DIR,
        'results_dir': RESULTS_DIR
    })

# ==================== STATIC FILES ====================

@app.route('/images/<path:filename>')
def serve_image(filename):
    file_path = os.path.join(INPUT_IMAGES_DIR, filename)
    if os.path.exists(file_path):
        return send_file(file_path)
    return jsonify({'error': 'File khÃ´ng tá»“n táº¡i'}), 404

# ==================== API: ATTENDANCE ====================

# ÄÆ°á»ng dáº«n cho attendance

@app.route('/api/attendance/analyze', methods=['POST'])
def analyze_attendance():
    """PhÃ¢n tÃ­ch file cháº¥m cÃ´ng vÃ  tÃ¬m cÃ¡c báº£n ghi thiáº¿u"""
    try:
        from src.attendance_processor import AttendanceProcessor
        
        processor = AttendanceProcessor(CHAMCONG_DIR)
        processor.scan_all_files()
        missing = processor.get_missing_records()
        summary = processor.get_summary()
        
        return jsonify({
            'success': True,
            'summary': summary,
            'missing_records': missing
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/attendance/export', methods=['POST'])
def export_attendance():
    """Xuáº¥t file Word giáº£i trÃ¬nh vá»›i áº£nh"""
    try:
        from src.attendance_processor import AttendanceProcessor
        from src.word_exporter import WordExporter
        
        data = request.json or {}
        project_name = data.get('project_name', 'Chung cÆ° TÃ¢n Thuáº­n ÄÃ´ng')
        month = data.get('month', None)
        
        # Xá»­ lÃ½ cháº¥m cÃ´ng
        processor = AttendanceProcessor(CHAMCONG_DIR)
        processor.scan_all_files()
        missing = processor.get_missing_records()
        
        if not missing:
            return jsonify({
                'success': True,
                'message': 'KhÃ´ng cÃ³ báº£n ghi thiáº¿u cáº§n giáº£i trÃ¬nh',
                'output_file': None
            })
        
        # Xuáº¥t Word
        exporter = WordExporter(PORTRAIT_DIR, RESULTS_DIR)
        output_file = exporter.create_summary_document(missing, project_name, month)
        
        return jsonify({
            'success': True,
            'message': f'ÄÃ£ xuáº¥t {len(missing)} báº£n ghi thiáº¿u',
            'output_file': os.path.basename(output_file),
            'total_missing': len(missing)
        })
    except Exception as e:
        import traceback
        return jsonify({'success': False, 'error': str(e), 'trace': traceback.format_exc()}), 500

@app.route('/api/attendance/portraits')
def get_portrait_stats():
    """Thá»‘ng kÃª áº£nh chÃ¢n dung"""
    try:
        from src.word_exporter import WordExporter
        
        exporter = WordExporter(PORTRAIT_DIR, RESULTS_DIR)
        stats = exporter.get_portrait_stats()
        return jsonify({'success': True, **stats})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

# ==================== API: FULL ANALYSIS (NEW) ====================

# Face matcher instance (lazy loaded)
_face_matcher = None

def get_face_matcher():
    """Get or create face matcher instance"""
    global _face_matcher
    try:
        from src.face_matcher import FaceMatcher
        if _face_matcher is None:
            send_log("⏳ Dang khoi tao Face Matcher (DeepFace)...", "info")
            portrait_dir = resolve_portrait_dir(BASE_DIR)
            _face_matcher = FaceMatcher(portrait_dir, log_callback=send_log)
            send_log(
                f"✅ Face Matcher san sang. PortraitDir={portrait_dir} "
                f"(n={len(_face_matcher.portrait_cache)}, imgs={_count_images_in_dir(portrait_dir)})",
                "success"
            )
        # If cache empty, retry with scan-based directory
        if _face_matcher and len(_face_matcher.portrait_cache) == 0:
            alt_dir = resolve_portrait_dir_by_scan(BASE_DIR)
            if alt_dir:
                send_log(
                    f"🔁 Cache rong, thu lai PortraitDir={alt_dir} (imgs={_count_images_in_dir(alt_dir)})",
                    "warning"
                )
                _face_matcher = FaceMatcher(alt_dir, log_callback=send_log)
                send_log(
                    f"✅ Face Matcher san sang. PortraitDir={alt_dir} "
                    f"(n={len(_face_matcher.portrait_cache)})",
                    "success"
                )
    except Exception as e:
        send_log(f"❌ Loi khoi tao FaceMatcher: {e}", "error")
        import traceback
        traceback.print_exc()
        return None
    return _face_matcher

@app.route('/api/analyze-full', methods=['POST'])
def analyze_full():
    """PhÃ¢n tÃ­ch tá»•ng há»£p: tÃ¬m ngÃ y thiáº¿u + match áº£nh camera báº±ng nháº­n diá»‡n khuÃ´n máº·t"""
    try:
        from src.attendance_processor import AttendanceProcessor
        
        # Step 1: PhÃ¢n tÃ­ch cháº¥m cÃ´ng
        send_log("ðŸ“‚ Step 1: Äang phÃ¢n tÃ­ch file cháº¥m cÃ´ng...", "info")
        processor = AttendanceProcessor(CHAMCONG_DIR)
        processor.scan_all_files()
        missing_records = processor.get_missing_records()
        summary = processor.get_summary()
        send_log(f"ðŸ“‹ TÃ¬m tháº¥y {len(missing_records)} báº£n ghi thiáº¿u tá»« {summary.get('total_persons', 0)} ngÆ°á»i", "info")
        
        # Step 2: Khá»Ÿi táº¡o face matcher
        send_log("ðŸ”§ Step 2: Äang khá»Ÿi táº¡o Face Matcher...", "info")
        matcher = get_face_matcher()
        if matcher:
            send_log("âœ… Face Matcher Ä‘Ã£ sáºµn sÃ ng", "success")
        else:
            send_log("âš ï¸ Face Matcher khÃ´ng kháº£ dá»¥ng, sáº½ dÃ¹ng fallback", "warning")
        
        # Step 3: Match áº£nh camera cho má»—i báº£n ghi thiáº¿u
        send_log(f"ðŸ” Step 3: Báº¯t Ä‘áº§u matching áº£nh cho {len(missing_records)} báº£n ghi...", "info")
        matched_count = 0
        
        for i, record in enumerate(missing_records):
            date_str = record['date']  # format: dd/mm/yyyy
            day = date_str.split('/')[0].zfill(2)  # extract dd
            person_name = record['person_name']
            
            # TÃ¬m thÆ° má»¥c ngÃ y tÆ°Æ¡ng á»©ng
            day_folder = os.path.join(INPUT_IMAGES_DIR, day)
            
            record['matched_image'] = None
            
            if os.path.exists(day_folder):
                images = get_image_files(day_folder)
                send_log(f"  [{i+1}/{len(missing_records)}] {person_name} (ngÃ y {day}): TÃ¬m tháº¥y {len(images)} áº£nh trong thÆ° má»¥c", "default")
                
                if images and matcher:
                    # DÃ¹ng face recognition Ä‘á»ƒ tÃ¬m áº£nh match
                    try:
                        matched_image = matcher.match_face_in_images(person_name, images)
                        if matched_image:
                            record['matched_image'] = matched_image
                            matched_count += 1
                            send_log(f"  [{i+1}/{len(missing_records)}] âœ“ {person_name} -> {os.path.basename(matched_image)}", "success")
                        else:
                            send_log(f"  [{i+1}/{len(missing_records)}] âŒ {person_name}: KhÃ´ng tÃ¬m tháº¥y áº£nh match", "warning")
                    except Exception as match_err:
                        send_log(f"  [{i+1}/{len(missing_records)}] âŒ {person_name}: Lá»—i matcher ({match_err})", "error")
                elif images:
                    send_log(f"  [{i+1}/{len(missing_records)}] âš ï¸ {person_name}: FaceMatcher chÆ°a sáºµn sÃ ng, bá» qua", "warning")
                else:
                    send_log(f"  [{i+1}/{len(missing_records)}] âš ï¸ ThÆ° má»¥c {day} rá»—ng", "warning")
            else:
                send_log(f"  [{i+1}/{len(missing_records)}] âŒ KhÃ´ng tÃ¬m tháº¥y thÆ° má»¥c: {day_folder}", "error")
                record['matched_image'] = None
        
        summary['total_matched'] = matched_count
        send_log(f"ðŸŽ‰ HoÃ n thÃ nh! Matched {matched_count}/{len(missing_records)} báº£n ghi", "success")
        
        return jsonify({
            'success': True,
            'summary': summary,
            'records': missing_records
        })
    except Exception as e:
        import traceback
        send_log(f"âŒ Lá»—i: {e}", "error")
        return jsonify({'success': False, 'error': str(e), 'trace': traceback.format_exc()}), 500

@app.route('/matched-image/<path:filepath>')
def serve_matched_image(filepath):
    """Serve áº£nh Ä‘Ã£ match"""
    # Decode URL path náº¿u cáº§n
    import urllib.parse
    filepath = urllib.parse.unquote(filepath)
    
    # Thá»­ vá»›i Ä‘Æ°á»ng dáº«n nguyÃªn gá»‘c
    if os.path.exists(filepath):
        return send_file(filepath)
    
    # Thá»­ vá»›i Ä‘Æ°á»ng dáº«n tuyá»‡t Ä‘á»‘i tá»« BASE_DIR
    abs_path = os.path.join(BASE_DIR, filepath)
    if os.path.exists(abs_path):
        return send_file(abs_path)
    
    # Thá»­ thay tháº¿ backslash/forward slash
    filepath_fixed = filepath.replace('/', os.sep).replace('\\', os.sep)
    abs_path_fixed = os.path.join(BASE_DIR, filepath_fixed)
    if os.path.exists(abs_path_fixed):
        return send_file(abs_path_fixed)
    
    print(f"[serve_matched_image] File khÃ´ng tá»“n táº¡i:")
    print(f"  filepath: {filepath}")
    print(f"  abs_path: {abs_path}")
    print(f"  abs_path_fixed: {abs_path_fixed}")
    print(f"  BASE_DIR: {BASE_DIR}")
    
    return jsonify({'error': 'File khÃ´ng tá»“n táº¡i', 'filepath': filepath}), 404

@app.route('/api/export-word', methods=['POST'])
def export_word():
    """Xuáº¥t file Word vá»›i áº£nh camera Ä‘Ã£ match"""
    try:
        from docx import Document
        from docx.shared import Inches, Cm
        
        data = request.json or {}
        project_name = data.get('project_name', 'Chung cÆ° TÃ¢n Thuáº­n ÄÃ´ng')
        month = data.get('month', '')
        records = data.get('records', [])
        
        if not records:
            return jsonify({'success': False, 'error': 'KhÃ´ng cÃ³ dá»¯ liá»‡u Ä‘á»ƒ xuáº¥t'})
        
        # Táº¡o document
        doc = Document()
        
        # TiÃªu Ä‘á»
        title = doc.add_paragraph()
        title.add_run(f'GIáº¢I TRÃŒNH CHáº¤M CÃ”NG - {project_name}').bold = True
        title.alignment = 1  # Center
        
        doc.add_paragraph(f'ThÃ¡ng: {month}')
        doc.add_paragraph()
        
        # Táº¡o báº£ng
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Header
        headers = ['TÃŠN', 'NGÃ€Y', 'GIáº¢I TRÃŒNH', 'HÃŒNH áº¢NH', 'GHI CHÃš']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        
        # ThÃªm dá»¯ liá»‡u
        for record in records:
            row = table.add_row()
            row.cells[0].text = record.get('person_name', '')
            row.cells[1].text = record.get('date', '')
            row.cells[2].text = record.get('issue_description', 'NhÃ¢n viÃªn cÃ³ trá»±c, bá»• sung')
            
            # ThÃªm áº£nh náº¿u cÃ³
            matched_image = record.get('matched_image')
            if matched_image and os.path.exists(matched_image):
                try:
                    run = row.cells[3].paragraphs[0].add_run()
                    run.add_picture(matched_image, width=Cm(3))
                except Exception:
                    row.cells[3].text = '[Lá»—i áº£nh]'
            else:
                row.cells[3].text = '[KhÃ´ng cÃ³ áº£nh]'
            
            row.cells[4].text = ''
        
        # LÆ°u file
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"GIAI_TRINH_{project_name.replace(' ', '_')}_{timestamp}.docx"
        output_path = os.path.join(RESULTS_DIR, filename)
        doc.save(output_path)
        
        return jsonify({
            'success': True,
            'filename': filename,
            'path': output_path
        })
    except Exception as e:
        import traceback
        return jsonify({'success': False, 'error': str(e), 'trace': traceback.format_exc()}), 500

# ==================== API: PDF EXTRACTION ====================

# Import PDF extractor
try:
    from src import pdf_extractor
    PDF_EXTRACTOR_AVAILABLE = True
except ImportError:
    PDF_EXTRACTOR_AVAILABLE = False

# ThÆ° má»¥c cho PDF
PDF_UPLOAD_DIR = os.path.join(BASE_DIR, "pdf_uploads")
PDF_OUTPUT_DIR = os.path.join(BASE_DIR, "pdf_extracted")
os.makedirs(PDF_UPLOAD_DIR, exist_ok=True)
os.makedirs(PDF_OUTPUT_DIR, exist_ok=True)

@app.route('/api/pdf/check')
def pdf_check_available():
    """Kiá»ƒm tra xem tÃ­nh nÄƒng PDF cÃ³ sáºµn khÃ´ng"""
    available = PDF_EXTRACTOR_AVAILABLE and pdf_extractor.is_available()
    return jsonify({
        'available': available,
        'message': 'Sáºµn sÃ ng' if available else 'Cáº§n cÃ i Ä‘áº·t: pip install pdf2docx PyMuPDF python-docx'
    })

@app.route('/api/pdf/upload', methods=['POST'])
def pdf_upload():
    """Upload file PDF"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'KhÃ´ng cÃ³ file Ä‘Æ°á»£c upload'}), 400
    
    file = request.files['file']
    if not file.filename:
        return jsonify({'success': False, 'error': 'TÃªn file khÃ´ng há»£p lá»‡'}), 400
    
    if not file.filename.lower().endswith('.pdf'):
        return jsonify({'success': False, 'error': 'Chá»‰ cháº¥p nháº­n file PDF'}), 400
    
    filename = secure_filename(file.filename)
    filepath = os.path.join(PDF_UPLOAD_DIR, filename)
    file.save(filepath)
    
    return jsonify({
        'success': True,
        'filename': filename,
        'filepath': filepath,
        'size': os.path.getsize(filepath)
    })

@app.route('/api/pdf/extract', methods=['POST'])
def pdf_extract():
    """Báº¯t Ä‘áº§u tÃ¡ch PDF thÃ nh cÃ¡c file Word"""
    if not PDF_EXTRACTOR_AVAILABLE or not pdf_extractor.is_available():
        return jsonify({
            'success': False, 
            'error': 'PDF Extractor khÃ´ng kháº£ dá»¥ng. Cáº§n cÃ i Ä‘áº·t: pip install pdf2docx PyMuPDF python-docx'
        }), 400
    
    data = request.json or {}
    filename = data.get('filename')
    
    if not filename:
        return jsonify({'success': False, 'error': 'Thiáº¿u tÃªn file'}), 400
    
    filepath = os.path.join(PDF_UPLOAD_DIR, secure_filename(filename))
    
    if not os.path.exists(filepath):
        return jsonify({'success': False, 'error': 'File PDF khÃ´ng tá»“n táº¡i'}), 404
    
    # Táº¡o thÆ° má»¥c output riÃªng cho file nÃ y
    base_name = os.path.splitext(filename)[0]
    output_dir = os.path.join(PDF_OUTPUT_DIR, base_name)
    
    # Báº¯t Ä‘áº§u task trong background
    task_id = pdf_extractor.start_extraction_task(filepath, output_dir)
    
    return jsonify({
        'success': True,
        'task_id': task_id,
        'message': 'ÄÃ£ báº¯t Ä‘áº§u tÃ¡ch PDF',
        'output_dir': output_dir
    })

@app.route('/api/pdf/status/<task_id>')
def pdf_status(task_id):
    """Kiá»ƒm tra tiáº¿n Ä‘á»™ tÃ¡ch PDF"""
    if not PDF_EXTRACTOR_AVAILABLE:
        return jsonify({'error': 'PDF Extractor khÃ´ng kháº£ dá»¥ng'}), 400
    
    task = pdf_extractor.get_task(task_id)
    if not task:
        return jsonify({'error': 'Task khÃ´ng tá»“n táº¡i'}), 404
    
    return jsonify(task.to_dict())

@app.route('/api/pdf/files')
def pdf_list_files():
    """Liá»‡t kÃª cÃ¡c file Word Ä‘Ã£ tÃ¡ch"""
    files = []
    
    if os.path.exists(PDF_OUTPUT_DIR):
        for folder in os.listdir(PDF_OUTPUT_DIR):
            folder_path = os.path.join(PDF_OUTPUT_DIR, folder)
            if os.path.isdir(folder_path):
                folder_files = pdf_extractor.list_extracted_files(folder_path) if PDF_EXTRACTOR_AVAILABLE else []
                files.append({
                    'folder': folder,
                    'path': folder_path,
                    'files': folder_files,
                    'count': len(folder_files)
                })
    
    return jsonify({
        'success': True,
        'folders': files,
        'output_dir': PDF_OUTPUT_DIR
    })

@app.route('/api/pdf/files/<folder>')
def pdf_list_folder_files(folder):
    """Liá»‡t kÃª cÃ¡c file Word trong má»™t thÆ° má»¥c"""
    folder_path = os.path.join(PDF_OUTPUT_DIR, secure_filename(folder))
    
    if not os.path.exists(folder_path):
        return jsonify({'error': 'ThÆ° má»¥c khÃ´ng tá»“n táº¡i'}), 404
    
    files = pdf_extractor.list_extracted_files(folder_path) if PDF_EXTRACTOR_AVAILABLE else []
    
    return jsonify({
        'success': True,
        'folder': folder,
        'files': files,
        'count': len(files)
    })

@app.route('/api/pdf/download/<folder>/<filename>')
def pdf_download(folder, filename):
    """Táº£i file Word Ä‘Ã£ tÃ¡ch"""
    # Decode URL-encoded names (khÃ´ng dÃ¹ng secure_filename vÃ¬ nÃ³ xÃ³a tiáº¿ng Viá»‡t)
    from urllib.parse import unquote
    folder = unquote(folder)
    filename = unquote(filename)
    
    # Báº£o vá»‡ path traversal
    if '..' in folder or '..' in filename or '/' in folder or '\\' in folder:
        return jsonify({'error': 'Invalid path'}), 400
    
    file_path = os.path.join(PDF_OUTPUT_DIR, folder, filename)
    
    # Kiá»ƒm tra file náº±m trong thÆ° má»¥c cho phÃ©p
    if not os.path.abspath(file_path).startswith(os.path.abspath(PDF_OUTPUT_DIR)):
        return jsonify({'error': 'Invalid path'}), 400
    
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    return jsonify({'error': 'File khÃ´ng tá»“n táº¡i', 'path': file_path}), 404

@app.route('/api/pdf/uploads')
def pdf_list_uploads():
    """Liá»‡t kÃª cÃ¡c file PDF Ä‘Ã£ upload"""
    files = []
    
    if os.path.exists(PDF_UPLOAD_DIR):
        for file in os.listdir(PDF_UPLOAD_DIR):
            if file.lower().endswith('.pdf'):
                filepath = os.path.join(PDF_UPLOAD_DIR, file)
                files.append({
                    'name': file,
                    'size': os.path.getsize(filepath),
                    'modified': datetime.fromtimestamp(os.path.getmtime(filepath)).isoformat()
                })
    
    files.sort(key=lambda x: x['modified'], reverse=True)
    
    return jsonify({
        'success': True,
        'files': files,
        'upload_dir': PDF_UPLOAD_DIR
    })


# ==================== API: EXCEL EXTRACTION ====================

# Thư mục cho Excel
EXCEL_UPLOAD_DIR = os.path.join(BASE_DIR, "excel_uploads")
EXCEL_PERSON_DIR = os.path.join(BASE_DIR, "excel_persons")
EXCEL_OUTPUT_DIR = os.path.join(BASE_DIR, "excel_extracted")  # Word chi tiet cham cong theo nguoi
EXCEL_FACE_OUTPUT_DIR = os.path.join(BASE_DIR, "excel_face_output")
os.makedirs(EXCEL_UPLOAD_DIR, exist_ok=True)
os.makedirs(EXCEL_PERSON_DIR, exist_ok=True)
os.makedirs(EXCEL_OUTPUT_DIR, exist_ok=True)
os.makedirs(EXCEL_FACE_OUTPUT_DIR, exist_ok=True)

# Task storage cho excel
excel_tasks = {}
excel_face_tasks = {}

class ExcelTask:
    def __init__(self, task_id):
        self.task_id = task_id
        self.status = 'pending'   # pending | running | completed | failed
        self.progress = 0
        self.total = 0
        self.current = ''
        self.files = []
        self.errors = []
        self.start_time = None
        self.end_time = None

    def to_dict(self):
        return {
            'task_id': self.task_id,
            'status': self.status,
            'progress': self.progress,
            'total': self.total,
            'current': self.current,
            'files': self.files,
            'errors': self.errors,
            'start_time': self.start_time.isoformat() if self.start_time else None,
            'end_time': self.end_time.isoformat() if self.end_time else None,
        }

class ExcelFaceTask:
    def __init__(self, task_id):
        self.task_id = task_id
        self.status = 'pending'
        self.progress = 0
        self.total = 0
        self.current = ''
        self.files = []
        self.errors = []
        self.start_time = None
        self.end_time = None

    def to_dict(self):
        return {
            'task_id': self.task_id,
            'status': self.status,
            'progress': self.progress,
            'total': self.total,
            'current': self.current,
            'files': self.files,
            'errors': self.errors,
            'start_time': self.start_time.isoformat() if self.start_time else None,
            'end_time': self.end_time.isoformat() if self.end_time else None,
        }

@app.route('/api/excel/upload', methods=['POST'])
def excel_upload():
    """Upload file Excel cháº¥m cÃ´ng (.xls/.xlsx)"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'KhÃ´ng cÃ³ file Ä‘Æ°á»£c upload'}), 400

    file = request.files['file']
    if not file.filename:
        return jsonify({'success': False, 'error': 'TÃªn file khÃ´ng há»£p lá»‡'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ('.xls', '.xlsx'):
        return jsonify({'success': False, 'error': 'Chá»‰ cháº¥p nháº­n file .xls hoáº·c .xlsx'}), 400

    # Giá»¯ tÃªn gá»‘c (cÃ³ tiáº¿ng Viá»‡t)
    filename = file.filename
    filepath = os.path.join(EXCEL_UPLOAD_DIR, filename)
    file.save(filepath)

    return jsonify({
        'success': True,
        'filename': filename,
        'filepath': filepath,
        'size': os.path.getsize(filepath)
    })

@app.route('/api/excel/extract', methods=['POST'])
def excel_extract():
    """Bat dau tach Excel -> file Excel theo nguoi + file Word de in"""
    try:
        data = request.json or {}
        filename = data.get('filename')
        if not filename:
            return jsonify({'success': False, 'error': 'Thiếu tên file'}), 400

        filepath = os.path.join(EXCEL_UPLOAD_DIR, filename)
        if not os.path.exists(filepath):
            return jsonify({'success': False, 'error': 'File Excel không tồn tại'}), 404

        # Tạo thư mục output riêng
        base_name = os.path.splitext(filename)[0]
        person_dir = os.path.join(EXCEL_PERSON_DIR, base_name)
        output_dir = os.path.join(EXCEL_OUTPUT_DIR, base_name)
        os.makedirs(person_dir, exist_ok=True)
        os.makedirs(output_dir, exist_ok=True)

        task_id = f"excel_{int(time.time() * 1000)}"
        task = ExcelTask(task_id)
        excel_tasks[task_id] = task

        def _run():
            task.status = 'running'
            task.start_time = datetime.now()
            try:
                send_log(f"📂 Đang đọc file Excel: {filename}", "info")
                from src.excel_splitter import ExcelAttendanceSplitter
                from src.excel_list_word_exporter import ExcelListWordExporter

                splitter = ExcelAttendanceSplitter(filepath)
                person_files, summaries = splitter.split(person_dir)
                task.total = len(summaries)
                send_log(f"✅ Tìm thấy {len(summaries)} nhân viên trong file", "success")

                exporter = ExcelListWordExporter(output_dir)

                for i, s in enumerate(summaries, 1):
                    task.current = s['name']
                    word_path = exporter.export_from_excel(person_files[i - 1])
                    if not word_path:
                        task.errors.append(f"Khong tao duoc Word cho {s['name']}")
                        send_log(f"⚠️ Không tạo được Word cho {s['name']}", "warning")
                    else:
                        task.files.append({
                            'name': os.path.basename(word_path),
                            'person': s['name'],
                            'days': s['rows'],
                            'present_rows': s['present_rows'],
                            'folder': base_name,
                        })
                    task.progress = i

                task.status = 'completed'
                send_log(
                    f"🎉 Hoàn tất! Đã tạo {len(task.files)} file Word in trong excel_extracted\\{base_name}",
                    "success"
                )
            except Exception as e:
                import traceback
                task.status = 'failed'
                task.errors.append(str(e))
                send_log(f"❌ Lỗi xử lý Excel: {e}", "error")
                traceback.print_exc()
            task.end_time = datetime.now()

        t = threading.Thread(target=_run, daemon=True)
        t.start()

        return jsonify({
            'success': True,
            'task_id': task_id,
            'message': f'Đã bắt đầu xử lý {filename}',
        })
    except Exception as e:
        import traceback
        return jsonify({'success': False, 'error': str(e), 'trace': traceback.format_exc()}), 500

@app.route('/api/excel/status/<task_id>')
def excel_status(task_id):
    """Kiá»ƒm tra tiáº¿n Ä‘á»™ tÃ¡ch Excel"""
    task = excel_tasks.get(task_id)
    if not task:
        return jsonify({'error': 'Task khÃ´ng tá»“n táº¡i'}), 404
    return jsonify(task.to_dict())

@app.route('/api/excel/files')
def excel_list_files():
    """Liệt kê các file Word chi tiết chấm công đã tạo từ Excel"""
    folders = []
    if os.path.exists(EXCEL_OUTPUT_DIR):
        for folder in os.listdir(EXCEL_OUTPUT_DIR):
            folder_path = os.path.join(EXCEL_OUTPUT_DIR, folder)
            if os.path.isdir(folder_path):
                word_files = [
                    {
                        'name': f,
                        'size': os.path.getsize(os.path.join(folder_path, f)),
                        'modified': datetime.fromtimestamp(
                            os.path.getmtime(os.path.join(folder_path, f))
                        ).isoformat()
                    }
                    for f in os.listdir(folder_path) if f.lower().endswith('.docx')
                ]
                word_files.sort(key=lambda x: x['name'])
                folders.append({
                    'folder': folder,
                    'files': word_files,
                    'count': len(word_files)
                })
    return jsonify({'success': True, 'folders': folders, 'output_dir': EXCEL_OUTPUT_DIR})

@app.route('/api/excel/download/<folder>/<filename>')
def excel_download(folder, filename):
    """Tải file Word chi tiết chấm công theo từng người"""
    from urllib.parse import unquote
    folder = unquote(folder)
    filename = unquote(filename)

    if '..' in folder or '..' in filename or '/' in folder or '\\' in folder:
        return jsonify({'error': 'Invalid path'}), 400

    file_path = os.path.join(EXCEL_OUTPUT_DIR, folder, filename)

    if not os.path.abspath(file_path).startswith(os.path.abspath(EXCEL_OUTPUT_DIR)):
        return jsonify({'error': 'Invalid path'}), 400

    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    return jsonify({'error': 'File khÃ´ng tá»“n táº¡i', 'path': file_path}), 404


@app.route('/api/excel/face/analyze', methods=['POST'])
def excel_face_analyze():
    """Phân tích khuôn mặt từ các file Excel đã tách"""
    try:
        data = request.json or {}
        folder = data.get('folder')
        threshold = data.get('distance_threshold')
        try:
            threshold = float(threshold) if threshold not in (None, '') else None
        except Exception:
            threshold = None
        if not folder:
            return jsonify({'success': False, 'error': 'Thiếu tên thư mục'}), 400

        input_dir = os.path.join(EXCEL_PERSON_DIR, folder)
        if not os.path.exists(input_dir):
            return jsonify({'success': False, 'error': 'Thư mục Excel đã tách không tồn tại'}), 404

        output_dir = os.path.join(EXCEL_FACE_OUTPUT_DIR, folder)
        os.makedirs(output_dir, exist_ok=True)

        task_id = f"excel_face_{int(time.time() * 1000)}"
        task = ExcelFaceTask(task_id)
        excel_face_tasks[task_id] = task

        def _run():
            task.status = 'running'
            task.start_time = datetime.now()
            try:
                send_log(f"🔍 Bắt đầu phân tích khuôn mặt cho thư mục: {folder}", "info")
                matcher = get_face_matcher()
                if matcher:
                    send_log("✅ Face Matcher đã sẵn sàng", "success")
                else:
                    send_log("⚠️ Face Matcher không khả dụng, sẽ bỏ qua tìm ảnh camera", "warning")

                from src.excel_face_analyzer import ExcelFaceAnalyzer
                analyzer = ExcelFaceAnalyzer(
                    PORTRAIT_DIR,
                    INPUT_IMAGES_DIR,
                    matcher,
                    accuracy_mode=True,
                    match_distance_threshold=threshold,
                    log_detail=True
                )

                def _log(msg, t='default'):
                    send_log(msg, t)

                files = analyzer.analyze_folder(input_dir, output_dir, log_callback=_log)
                task.total = len(files)
                for i, f in enumerate(files, 1):
                    task.current = os.path.basename(f)
                    task.files.append({
                        'name': os.path.basename(f),
                        'folder': folder,
                    })
                    task.progress = i

                task.status = 'completed'
                send_log(f"🎉 Hoàn tất! Đã xuất {len(files)} file Word", "success")
            except Exception as e:
                import traceback
                task.status = 'failed'
                task.errors.append(str(e))
                send_log(f"❌ Lỗi phân tích Excel: {e}", "error")
                traceback.print_exc()
            task.end_time = datetime.now()

        t = threading.Thread(target=_run, daemon=True)
        t.start()

        return jsonify({
            'success': True,
            'task_id': task_id,
            'message': f'Đã bắt đầu phân tích {folder}',
        })
    except Exception as e:
        import traceback
        return jsonify({'success': False, 'error': str(e), 'trace': traceback.format_exc()}), 500


@app.route('/api/excel/face/status/<task_id>')
def excel_face_status(task_id):
    """Kiểm tra tiến độ phân tích khuôn mặt từ Excel"""
    task = excel_face_tasks.get(task_id)
    if not task:
        return jsonify({'error': 'Task không tồn tại'}), 404
    return jsonify(task.to_dict())


@app.route('/api/excel/face/files')
def excel_face_files():
    """Liệt kê các file Word đã phân tích từ Excel"""
    folders = []
    if os.path.exists(EXCEL_FACE_OUTPUT_DIR):
        for folder in os.listdir(EXCEL_FACE_OUTPUT_DIR):
            folder_path = os.path.join(EXCEL_FACE_OUTPUT_DIR, folder)
            if os.path.isdir(folder_path):
                word_files = [
                    {
                        'name': f,
                        'size': os.path.getsize(os.path.join(folder_path, f)),
                        'modified': datetime.fromtimestamp(
                            os.path.getmtime(os.path.join(folder_path, f))
                        ).isoformat()
                    }
                    for f in os.listdir(folder_path) if f.lower().endswith('.docx')
                ]
                word_files.sort(key=lambda x: x['name'])
                folders.append({
                    'folder': folder,
                    'files': word_files,
                    'count': len(word_files)
                })
    return jsonify({'success': True, 'folders': folders, 'output_dir': EXCEL_FACE_OUTPUT_DIR})


@app.route('/api/excel/face/download/<folder>/<filename>')
def excel_face_download(folder, filename):
    """Tải file Word đã phân tích từ Excel"""
    from urllib.parse import unquote
    folder = unquote(folder)
    filename = unquote(filename)

    if '..' in folder or '..' in filename or '/' in folder or '\\' in folder:
        return jsonify({'error': 'Invalid path'}), 400

    file_path = os.path.join(EXCEL_FACE_OUTPUT_DIR, folder, filename)

    if not os.path.abspath(file_path).startswith(os.path.abspath(EXCEL_FACE_OUTPUT_DIR)):
        return jsonify({'error': 'Invalid path'}), 400

    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    return jsonify({'error': 'File không tồn tại', 'path': file_path}), 404

@app.route('/api/excel/uploads')
def excel_list_uploads():
    """Liá»‡t kÃª cÃ¡c file Excel Ä‘Ã£ upload"""
    files = []
    if os.path.exists(EXCEL_UPLOAD_DIR):
        for f in os.listdir(EXCEL_UPLOAD_DIR):
            ext = os.path.splitext(f)[1].lower()
            if ext in ('.xls', '.xlsx'):
                fp = os.path.join(EXCEL_UPLOAD_DIR, f)
                files.append({
                    'name': f,
                    'size': os.path.getsize(fp),
                    'modified': datetime.fromtimestamp(os.path.getmtime(fp)).isoformat()
                })
    files.sort(key=lambda x: x['modified'], reverse=True)
    return jsonify({'success': True, 'files': files, 'upload_dir': EXCEL_UPLOAD_DIR})

# ==================== MAIN ====================


if __name__ == '__main__':
    print(f"""
â•”â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•—
â•‘     PHáº¦N Má»€M NHáº¬N DIá»†N KHUÃ”N Máº¶T & TRÃCH XUáº¤T NGÃ€Y THÃNG    â•‘
â• â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•£
â•‘  Server Ä‘ang cháº¡y táº¡i: http://localhost:{FLASK_PORT}                   â•‘
â•‘  Input Images: {INPUT_IMAGES_DIR:<44} â•‘
â•‘  Database:     {DATABASE_DIR:<44} â•‘
â•‘  Results:      {RESULTS_DIR:<44} â•‘
â•šâ•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    """)
    
    print("Äang quÃ©t database...")
    scan_database()
    print(f"ÄÃ£ load {len(database)} ngÆ°á»i trong database\n")
    
    app.run(host=FLASK_HOST, port=FLASK_PORT, debug=FLASK_DEBUG, threaded=True)
