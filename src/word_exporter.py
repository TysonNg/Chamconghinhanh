# -*- coding: utf-8 -*-
"""
Module xuất file Word giải trình với ảnh chân dung
"""

import os
import re
import unicodedata
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import List, Dict, Optional


class WordExporter:
    """Xuất file Word giải trình theo mẫu có sẵn, với ảnh chân dung"""
    
    def __init__(self, portrait_dir: str, output_dir: str):
        self.portrait_dir = portrait_dir  # Thư mục Ảnh BV (có subfolder theo tên)
        self.output_dir = output_dir
        self.portrait_cache = {}  # Cache mapping tên -> danh sách ảnh
        self._scan_portraits()
    
    def _scan_portraits(self):
        """Quét thư mục ảnh chân dung, hỗ trợ cả file trực tiếp và subfolder"""
        if not os.path.exists(self.portrait_dir):
            print(f"Thư mục ảnh không tồn tại: {self.portrait_dir}")
            return
        
        supported_ext = {'.jpg', '.jpeg', '.png', '.bmp'}
        
        for item in os.listdir(self.portrait_dir):
            item_path = os.path.join(self.portrait_dir, item)
            
            if os.path.isdir(item_path):
                # Subfolder theo tên người
                person_name = item
                images = []
                for file in os.listdir(item_path):
                    ext = os.path.splitext(file)[1].lower()
                    if ext in supported_ext:
                        images.append(os.path.join(item_path, file))
                
                if images:
                    normalized_name = self._normalize_name(person_name)
                    self.portrait_cache[normalized_name] = images
            else:
                # File ảnh trực tiếp (cấu trúc cũ)
                ext = os.path.splitext(item)[1].lower()
                if ext in supported_ext:
                    person_name = os.path.splitext(item)[0]
                    normalized_name = self._normalize_name(person_name)
                    if normalized_name not in self.portrait_cache:
                        self.portrait_cache[normalized_name] = []
                    self.portrait_cache[normalized_name].append(item_path)
        
        print(f"Đã tìm thấy ảnh của {len(self.portrait_cache)} người")
    
    def _normalize_name(self, name: str) -> str:
        """Chuẩn hóa tên để so sánh (bỏ dấu, lowercase, bỏ khoảng trắng thừa)"""
        # Bỏ dấu tiếng Việt
        name = unicodedata.normalize('NFD', name)
        name = ''.join(c for c in name if unicodedata.category(c) != 'Mn')
        # Lowercase và xử lý khoảng trắng
        name = name.lower().strip()
        name = re.sub(r'\s+', ' ', name)
        return name
    
    def find_portrait(self, person_name: str) -> Optional[List[str]]:
        """Tìm ảnh chân dung theo tên (fuzzy match)"""
        normalized_search = self._normalize_name(person_name)
        
        # Tìm chính xác
        if normalized_search in self.portrait_cache:
            return self.portrait_cache[normalized_search]
        
        # Fuzzy match - tìm tên chứa hoặc được chứa
        for cached_name, images in self.portrait_cache.items():
            if normalized_search in cached_name or cached_name in normalized_search:
                return images
        
        # So sánh từng từ
        search_words = set(normalized_search.split())
        best_match = None
        best_score = 0
        
        for cached_name, images in self.portrait_cache.items():
            cached_words = set(cached_name.split())
            common = len(search_words & cached_words)
            if common > best_score:
                best_score = common
                best_match = images
        
        if best_score >= 2:  # Ít nhất 2 từ trùng
            return best_match
        
        return None
    
    def create_summary_document(self, missing_records: List[Dict], 
                                 project_name: str = "Chung cư Tân Thuận Đông",
                                 month: str = None) -> str:
        """Tạo file Word giải trình theo format mẫu"""
        
        if month is None:
            month = datetime.now().strftime("%m/%Y")
        
        doc = Document()
        
        # === Header ===
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_para.add_run("Kính gửi: Ban Quản Lý Khu Đô Thị")
        header_run.bold = True
        
        # Tiêu đề
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("GIẢI TRÌNH CÔNG NHÂN VIÊN")
        title_run.bold = True
        title_run.font.size = Pt(14)
        
        # Thời gian
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(f"Tháng {month}")
        
        # Nội dung
        content_para = doc.add_paragraph()
        content_para.add_run(f"(V/v: nhân viên bảo vệ Dự án: {project_name} thiếu dữ liệu chấm công vân tay)")
        
        doc.add_paragraph()  # Dòng trống
        
        # === Bảng chính ===
        # Tạo bảng với 5 cột: TÊN, NGÀY, GIẢI TRÌNH, HÌNH ẢNH THỰC TẾ, GHI CHÚ
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ['TÊN', 'NGÀY', 'GIẢI TRÌNH', 'HÌNH ẢNH THỰC TẾ', 'GHI CHÚ']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Đặt độ rộng cột
        widths = [Cm(3.5), Cm(2.5), Cm(4), Cm(5), Cm(2.5)]
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = width
        
        # Điền dữ liệu
        current_person = None
        for record in missing_records:
            row = table.add_row()
            cells = row.cells
            
            # Tên (chỉ hiện khi đổi người)
            if record['person_name'] != current_person:
                cells[0].text = record['person_name']
                current_person = record['person_name']
            
            # Ngày
            cells[1].text = record['date']
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Giải trình
            cells[2].text = record.get('explanation', 'Nhân viên có trực, bổ sung')
            
            # Hình ảnh - tìm và chèn ảnh
            portraits = self.find_portrait(record['person_name'])
            if portraits:
                try:
                    # Chèn ảnh đầu tiên
                    paragraph = cells[3].paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(portraits[0], width=Cm(4))
                except Exception as e:
                    cells[3].text = f"[Lỗi ảnh: {str(e)[:30]}]"
            else:
                cells[3].text = "[Không tìm thấy ảnh]"
            
            # Ghi chú
            cells[4].text = ''
        
        # Lưu file
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f'GIAI_TRINH_{project_name.replace(" ", "_")}_{timestamp}.docx'
        output_path = os.path.join(self.output_dir, output_filename)
        
        os.makedirs(self.output_dir, exist_ok=True)
        doc.save(output_path)
        
        print(f"Đã tạo file: {output_path}")
        return output_path
    
    def get_portrait_stats(self) -> Dict:
        """Thống kê ảnh chân dung"""
        return {
            'total_persons': len(self.portrait_cache),
            'total_images': sum(len(imgs) for imgs in self.portrait_cache.values()),
            'persons': list(self.portrait_cache.keys())
        }


# Test
if __name__ == '__main__':
    import sys
    sys.stdout.reconfigure(encoding='utf-8')
    
    from attendance_processor import AttendanceProcessor
    
    # Paths
    portrait_dir = r'd:\Projects\phan mem quet mat\Ảnh BV'
    chamcong_dir = r'd:\Projects\phan mem quet mat\chamcong'
    output_dir = r'd:\Projects\phan mem quet mat\results'
    
    # Xử lý chấm công
    print("=== Quét chấm công ===")
    processor = AttendanceProcessor(chamcong_dir)
    processor.scan_all_files()
    missing = processor.get_missing_records()
    print(f"Tìm thấy {len(missing)} bản ghi thiếu")
    
    # Xuất Word
    print("\n=== Xuất Word ===")
    exporter = WordExporter(portrait_dir, output_dir)
    
    print("\n--- Thống kê ảnh ---")
    stats = exporter.get_portrait_stats()
    print(f"Tổng số người có ảnh: {stats['total_persons']}")
    print(f"Tổng số ảnh: {stats['total_images']}")
    
    if missing:
        output_file = exporter.create_summary_document(
            missing,
            project_name="Chung cư Tân Thuận Đông",
            month="12/2025"
        )
        print(f"\nĐã xuất file: {output_file}")
    else:
        print("\nKhông có bản ghi thiếu để xuất")
