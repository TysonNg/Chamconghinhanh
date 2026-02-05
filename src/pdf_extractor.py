# -*- coding: utf-8 -*-
"""
PDF Extractor Module
Tách bảng từ file PDF thành các file Word riêng biệt
Tích hợp vào Flask app
"""

import os
import re
import threading
from datetime import datetime

# Thử import các thư viện cần thiết
try:
    from pdf2docx import Converter
    import fitz  # PyMuPDF
    from docx import Document
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False

# Configuration
PDF_OUTPUT_DIR = "pdf_extracted"


class PDFExtractorTask:
    """Task để track tiến độ tách PDF"""
    
    def __init__(self, task_id):
        self.task_id = task_id
        self.status = 'pending'  # pending, running, completed, error
        self.progress = 0
        self.total = 0
        self.current_page = 0
        self.message = ''
        self.output_dir = ''
        self.files_created = []
        self.error = None
        self.start_time = None
        self.end_time = None
    
    def to_dict(self):
        return {
            'task_id': self.task_id,
            'status': self.status,
            'progress': self.progress,
            'total': self.total,
            'current_page': self.current_page,
            'message': self.message,
            'output_dir': self.output_dir,
            'files_created': self.files_created,
            'error': self.error,
            'start_time': self.start_time.isoformat() if self.start_time else None,
            'end_time': self.end_time.isoformat() if self.end_time else None
        }


# Global task storage
pdf_tasks = {}


def get_employee_name_from_docx(docx_path):
    """Extract employee name from converted docx file"""
    try:
        doc = Document(docx_path)
        
        for para in doc.paragraphs:
            text = para.text
            # Look for employee name pattern
            name_match = re.search(r'Tên\s*nhân viên[:\s]*([^\n\d]+?)(?:Phòng|$)', text)
            if name_match:
                name = name_match.group(1).strip()
                name = re.sub(r'[<>:"/\\|?*]', '', name)
                if name and len(name) > 2:
                    return name
            
            # Alternative pattern
            name_match = re.search(r'Mã\s*nhân viên[:\s]*\d+\s*Tên\s*nhân viên[:\s]*([^\n]+?)(?:Phòng|$)', text)
            if name_match:
                name = name_match.group(1).strip()
                name = re.sub(r'[<>:"/\\|?*]', '', name)
                if name and len(name) > 2:
                    return name
        
        # Try tables
        for table in doc.tables:
            for row in table.rows[:5]:
                for cell in row.cells:
                    text = cell.text
                    name_match = re.search(r'Tên\s*nhân viên[:\s]*([^\n]+?)(?:Phòng|$)', text)
                    if name_match:
                        name = name_match.group(1).strip()
                        name = re.sub(r'[<>:"/\\|?*]', '', name)
                        if name and len(name) > 2:
                            return name
    except Exception as e:
        pass
    
    return None


def extract_pdf_to_word(pdf_path, output_dir, task):
    """
    Chuyển PDF sang nhiều file Word, mỗi trang là 1 file
    
    Args:
        pdf_path: Đường dẫn file PDF
        output_dir: Thư mục xuất file Word
        task: PDFExtractorTask để track tiến độ
    
    Returns:
        List các file Word đã tạo
    """
    if not PDF2DOCX_AVAILABLE:
        raise ImportError("Cần cài đặt: pip install pdf2docx PyMuPDF python-docx")
    
    # Create output directory
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    task.output_dir = output_dir
    task.status = 'running'
    task.start_time = datetime.now()
    task.message = 'Đang đọc file PDF...'
    
    try:
        # Get page count
        pdf_doc = fitz.open(pdf_path)
        page_count = len(pdf_doc)
        pdf_doc.close()
        
        task.total = page_count
        task.message = f'PDF có {page_count} trang'
        
        files_created = []
        
        # Convert each page
        for page_num in range(page_count):
            task.current_page = page_num + 1
            task.progress = int((page_num / page_count) * 100)
            task.message = f'Đang xử lý trang {page_num + 1}/{page_count}...'
            
            # Temporary filename
            temp_file = os.path.join(output_dir, f"_temp_page_{page_num + 1}.docx")
            
            # Convert single page
            cv = Converter(pdf_path)
            cv.convert(temp_file, start=page_num, end=page_num + 1)
            cv.close()
            
            # Try to extract employee name
            name = get_employee_name_from_docx(temp_file)
            
            if name and len(name) > 2:
                final_name = name[:50]  # Limit length
            else:
                final_name = f"Page_{page_num + 1:02d}"
            
            # Clean filename
            final_name = re.sub(r'[<>:"/\\|?*]', '', final_name).strip()
            
            # Final filename
            final_file = os.path.join(output_dir, f"{final_name}.docx")
            
            # Handle duplicates
            counter = 1
            base_name = final_name
            while os.path.exists(final_file) and final_file != temp_file:
                final_name = f"{base_name}_{counter}"
                final_file = os.path.join(output_dir, f"{final_name}.docx")
                counter += 1
            
            # Rename temp file to final name
            if temp_file != final_file:
                if os.path.exists(final_file):
                    os.remove(final_file)
                os.rename(temp_file, final_file)
            
            files_created.append({
                'name': f"{final_name}.docx",
                'path': final_file,
                'page': page_num + 1
            })
        
        task.files_created = files_created
        task.progress = 100
        task.status = 'completed'
        task.message = f'Hoàn thành! Đã tạo {len(files_created)} file Word.'
        task.end_time = datetime.now()
        
        return files_created
        
    except Exception as e:
        task.status = 'error'
        task.error = str(e)
        task.message = f'Lỗi: {str(e)}'
        task.end_time = datetime.now()
        raise


def start_extraction_task(pdf_path, output_dir):
    """
    Bắt đầu task tách PDF trong background thread
    
    Returns:
        task_id
    """
    import uuid
    task_id = str(uuid.uuid4())[:8]
    
    task = PDFExtractorTask(task_id)
    pdf_tasks[task_id] = task
    
    def run():
        try:
            extract_pdf_to_word(pdf_path, output_dir, task)
        except Exception as e:
            task.status = 'error'
            task.error = str(e)
    
    thread = threading.Thread(target=run, daemon=True)
    thread.start()
    
    return task_id


def get_task(task_id):
    """Lấy thông tin task theo ID"""
    return pdf_tasks.get(task_id)


def list_extracted_files(output_dir):
    """Liệt kê các file Word đã tách"""
    if not os.path.exists(output_dir):
        return []
    
    files = []
    for filename in os.listdir(output_dir):
        if filename.endswith('.docx') and not filename.startswith('_temp'):
            filepath = os.path.join(output_dir, filename)
            files.append({
                'name': filename,
                'path': filepath,
                'size': os.path.getsize(filepath),
                'modified': datetime.fromtimestamp(os.path.getmtime(filepath)).isoformat()
            })
    
    # Sort by name
    files.sort(key=lambda x: x['name'])
    return files


def is_available():
    """Kiểm tra xem module có sẵn sàng không"""
    return PDF2DOCX_AVAILABLE
