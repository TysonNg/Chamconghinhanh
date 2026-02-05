# -*- coding: utf-8 -*-
"""
Script to convert PDF pages to Word documents using pdf2docx
Each page becomes a separate Word file
"""

import os
import re
from pdf2docx import Converter

# Configuration
PDF_PATH = r"CHẤM CÔNG BẢO VỆ T1.2026.pdf"
OUTPUT_DIR = r"chamcong_t1_2026_final"

def get_employee_name_from_docx(docx_path):
    """Extract employee name from converted docx file"""
    try:
        from docx import Document
        doc = Document(docx_path)
        
        for para in doc.paragraphs:
            text = para.text
            # Look for employee name pattern
            name_match = re.search(r'Tên\s*nhân viên[:\s]*([^\n\d]+?)(?:Phòng|$)', text)
            if name_match:
                name = name_match.group(1).strip()
                name = re.sub(r'[<>:"/\\|?*]', '', name)
                if name and len(name) > 2:
                    return name
            
            # Alternative pattern
            name_match = re.search(r'Mã\s*nhân viên[:\s]*\d+\s*Tên\s*nhân viên[:\s]*([^\n]+?)(?:Phòng|$)', text)
            if name_match:
                name = name_match.group(1).strip()
                name = re.sub(r'[<>:"/\\|?*]', '', name)
                if name and len(name) > 2:
                    return name
        
        # Try tables
        for table in doc.tables:
            for row in table.rows[:5]:
                for cell in row.cells:
                    text = cell.text
                    name_match = re.search(r'Tên\s*nhân viên[:\s]*([^\n]+?)(?:Phòng|$)', text)
                    if name_match:
                        name = name_match.group(1).strip()
                        name = re.sub(r'[<>:"/\\|?*]', '', name)
                        if name and len(name) > 2:
                            return name
    except Exception as e:
        print(f"  Warning: Could not extract name: {e}")
    
    return None

def main():
    print("="*60)
    print("PDF to Word Conversion Tool (pdf2docx)")
    print("="*60)
    
    # Create output directory
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
    print(f"Output directory: {OUTPUT_DIR}")
    
    print(f"\nConverting PDF: {PDF_PATH}")
    
    try:
        # Get page count using PyMuPDF
        import fitz
        pdf_doc = fitz.open(PDF_PATH)
        page_count = len(pdf_doc)
        pdf_doc.close()
        print(f"PDF has {page_count} pages")
        
        # Convert each page to separate Word file
        for page_num in range(page_count):
            print(f"Converting page {page_num + 1}/{page_count}...", end=" ")
            
            # Temporary filename
            temp_file = os.path.join(OUTPUT_DIR, f"_temp_page_{page_num + 1}.docx")
            
            # Convert single page
            cv = Converter(PDF_PATH)
            cv.convert(temp_file, start=page_num, end=page_num + 1)
            cv.close()
            
            # Try to extract employee name
            name = get_employee_name_from_docx(temp_file)
            
            if name and len(name) > 2:
                final_name = name[:50]  # Limit length
            else:
                final_name = f"Page_{page_num + 1:02d}"
            
            # Clean filename
            final_name = re.sub(r'[<>:"/\\|?*]', '', final_name).strip()
            
            # Final filename
            final_file = os.path.join(OUTPUT_DIR, f"{final_name}.docx")
            
            # Handle duplicates
            counter = 1
            base_name = final_name
            while os.path.exists(final_file) and final_file != temp_file:
                final_name = f"{base_name}_{counter}"
                final_file = os.path.join(OUTPUT_DIR, f"{final_name}.docx")
                counter += 1
            
            # Rename temp file to final name
            if temp_file != final_file:
                os.rename(temp_file, final_file)
            
            print(f"Created: {final_name}.docx")
        
        print(f"\n{'='*60}")
        files_created = len([f for f in os.listdir(OUTPUT_DIR) if f.endswith('.docx')])
        print(f"Done! Created {files_created} Word documents in '{OUTPUT_DIR}'")
        print(f"{'='*60}")
        
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
